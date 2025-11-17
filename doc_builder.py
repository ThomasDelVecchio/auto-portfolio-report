import os
import sys
import shutil

import numpy as np
import pandas as pd
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

from time_utils import get_eastern_now
from config import TARGET_PORTFOLIO_VALUE, monthly_contrib, EXTRA_OUTPUT_DIRS
from helpers import fmt_pct, fmt_pct_level, fmt_dollar, weighted_avg


def build_report(
    df,
    asset_df,
    asset_targets_df,
    returns_df,
    dollar_pl_df,
    bench_df,
    total_value,
    summary_1d_return,
    total_1d_pl,
    port_mtd,
    total_mtd_pl,
    port_ytd,
    total_ytd_pl,
    sector_stream,
    ticker_pie_stream,
    asset_pie_stream,
    growth_stream,
    compound_stream,
    vol_stream,
    risk_stream,
):
    # ---- Executive Summary Metrics ----

    summary_total_value = total_value
    summary_target_value = TARGET_PORTFOLIO_VALUE
    summary_num_holdings = len(df)

    # NOTE:
    # summary_1d_return, total_1d_pl, total_mtd_pl, total_ytd_pl
    # are computed in the TWR / benchmark section using portfolio_values
    # (portfolio value series). They are TRUE time-weighted returns and
    # fully consistent with the dollar P/L.

    # Join returns and $ P/L for ranking
    ret_pl = returns_df.merge(dollar_pl_df, on="Ticker", how="left")

    top_1m_line = "N/A"
    bottom_1m_line = "N/A"
    best_1d_line = None  # optional

    if not ret_pl.empty:
        # Top & bottom 1M performers
        valid_1m = ret_pl.dropna(subset=["1M %"])
        if not valid_1m.empty:
            top_1m_row = valid_1m.sort_values("1M %", ascending=False).iloc[0]
            top_1m_line = (
                f"{top_1m_row['Ticker']} "
                f"({fmt_pct(top_1m_row['1M %'])}, "
                f"{fmt_dollar(top_1m_row.get('1M $'))})"
            )

            bottom_1m_row = valid_1m.sort_values("1M %", ascending=True).iloc[0]
            bottom_1m_line = (
                f"{bottom_1m_row['Ticker']} "
                f"({fmt_pct(bottom_1m_row['1M %'])}, "
                f"{fmt_dollar(bottom_1m_row.get('1M $'))})"
            )

        # Best 1D performer (optional)
        valid_1d = ret_pl.dropna(subset=["1D %"])
        if not valid_1d.empty:
            best_1d_row = valid_1d.sort_values("1D %", ascending=False).iloc[0]
            best_1d_line = (
                f"{best_1d_row['Ticker']} "
                f"({fmt_pct(best_1d_row['1D %'])}, "
                f"{fmt_dollar(best_1d_row.get('1D $'))})"
            )

    # Top 3 holdings concentration
    if summary_num_holdings > 0:
        top3_pct = float(df["allocation_pct"].nlargest(min(3, summary_num_holdings)).sum())
    else:
        top3_pct = float("nan")

    # Largest asset class by actual allocation
    if not asset_df.empty:
        largest_ac_row = asset_df.sort_values("allocation_pct", ascending=False).iloc[0]
        largest_ac_name = largest_ac_row["asset_class"]
        largest_ac_pct = float(largest_ac_row["allocation_pct"])
    else:
        largest_ac_name = "N/A"
        largest_ac_pct = float("nan")

    # Overweight / underweight asset classes vs targets
    largest_overweight_str = "N/A"
    largest_underweight_str = "N/A"

    if asset_targets_df is not None and not asset_targets_df.empty:
        ac_compare_for_summary = asset_df.merge(
            asset_targets_df, on="asset_class", how="left", suffixes=("", "_target")
        )
        ac_compare_for_summary["target_pct"] = ac_compare_for_summary["target_pct"].fillna(0.0)
        ac_compare_for_summary["delta_pct"] = (
            ac_compare_for_summary["allocation_pct"] - ac_compare_for_summary["target_pct"]
        )

        ow = ac_compare_for_summary[ac_compare_for_summary["delta_pct"] > 0]
        uw = ac_compare_for_summary[ac_compare_for_summary["delta_pct"] < 0]

        if not ow.empty:
            ow_row = ow.sort_values("delta_pct", ascending=False).iloc[0]
            largest_overweight_str = (
                f"{ow_row['asset_class']} "
                f"({ow_row['allocation_pct']:.2f}% vs {ow_row['target_pct']:.2f}%)"
            )

        if not uw.empty:
            uw_row = uw.sort_values("delta_pct", ascending=True).iloc[0]
            largest_underweight_str = (
                f"{uw_row['asset_class']} "
                f"({uw_row['allocation_pct']:.2f}% vs {uw_row['target_pct']:.2f}%)"
            )

    # Rebalancing need & % of tickers on target (within ±5% band)
    rebalance_need = float(df["contribute_to_target"].fillna(0).sum()) if not df.empty else 0.0

    pct_tickers_on_target = None
    if "target_pct" in df.columns and not df["target_pct"].isna().all():
        has_target = df["target_pct"].fillna(0) > 0
        n_with_target = int(has_target.sum())
        if n_with_target > 0:
            drift = (df.loc[has_target, "allocation_pct"] - df.loc[has_target, "target_pct"]).abs()
            within_band = drift <= 5.0  # within ±5% band
            pct_tickers_on_target = float(within_band.mean() * 100.0)

    # Benchmark positioning vs S&P 500 (excess returns MTD/YTD)
    vs_sp500_mtd_str = "N/A"
    vs_sp500_ytd_str = "N/A"

    try:
        bench_df["MTD %"] = pd.to_numeric(bench_df["MTD %"], errors="coerce")
        bench_df["YTD %"] = pd.to_numeric(bench_df["YTD %"], errors="coerce")

        if (
            "Portfolio (Live)" in bench_df["Benchmark"].values
            and "S&P 500" in bench_df["Benchmark"].values
        ):
            port_row_b = bench_df.loc[bench_df["Benchmark"] == "Portfolio (Live)"].iloc[0]
            sp_row = bench_df.loc[bench_df["Benchmark"] == "S&P 500"].iloc[0]

            if not pd.isna(port_row_b["MTD %"]) and not pd.isna(sp_row["MTD %"]):
                vs_sp500_mtd = float(port_row_b["MTD %"] - sp_row["MTD %"])
                vs_sp500_mtd_str = fmt_pct(vs_sp500_mtd)

            if not pd.isna(port_row_b["YTD %"]) and not pd.isna(sp_row["YTD %"]):
                vs_sp500_ytd = float(port_row_b["YTD %"] - sp_row["YTD %"])
                vs_sp500_ytd_str = fmt_pct(vs_sp500_ytd)
    except Exception:
        pass

    # Now create the document
    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style._element.rPr.rFonts.set(qn("w:eastAsia"), "Calibri")
    style.font.size = Pt(11)

    def add_table(headers, rows, right_align_cols=None):
        table = doc.add_table(rows=1, cols=len(headers))
        table.style = "Light Grid Accent 1"

        hdr_cells = table.rows[0].cells
        for i, h in enumerate(headers):
            hdr_cells[i].text = h
            for p in hdr_cells[i].paragraphs:
                for r in p.runs:
                    r.bold = True

        for row_data in rows:
            row_cells = table.add_row().cells
            for i, val in enumerate(row_data):
                row_cells[i].text = str(val)

        # Repeat header row on each page
        hdr_row = table.rows[0]
        tr = hdr_row._tr
        trPr = tr.get_or_add_trPr()
        tbl_header = OxmlElement("w:tblHeader")
        trPr.append(tbl_header)

        # Prevent splitting rows; keep rows together
        row_count = len(table.rows)
        for idx, row in enumerate(table.rows):
            tr = row._tr
            trPr = tr.get_or_add_trPr()

            cant_split = OxmlElement("w:cantSplit")
            trPr.append(cant_split)

            for cell in row.cells:
                for p in cell.paragraphs:
                    pf = p.paragraph_format
                    pf.keep_together = True
                    pf.keep_with_next = (idx < row_count - 1)

        # Right-align numeric columns
        if right_align_cols is not None:
            for row in table.rows[1:]:
                for col_idx in right_align_cols:
                    if col_idx < len(row.cells):
                        cell = row.cells[col_idx]
                        for p in cell.paragraphs:
                            p.alignment = WD_ALIGN_PARAGRAPH.RIGHT

        return table

    # Cover
    doc.add_paragraph("\n\n\n\n\n")

    cover = doc.add_paragraph("Comprehensive Investment Report — Live Portfolio")
    cover.runs[0].bold = True
    cover.runs[0].font.size = Pt(26)
    cover.alignment = WD_ALIGN_PARAGRAPH.CENTER

    subtitle = doc.add_paragraph(
        "Automated Summary, Diversification, and Performance vs Benchmarks"
    )
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.runs[0].font.size = Pt(14)
    subtitle.runs[0].font.color.rgb = RGBColor(70, 130, 180)

    timestamp_str = get_eastern_now().strftime("%A, %B %d, %Y — %I:%M %p ET")
    ts = doc.add_paragraph(f"Report Run: {timestamp_str}")
    ts.alignment = WD_ALIGN_PARAGRAPH.CENTER
    ts.runs[0].font.size = Pt(12)
    ts.runs[0].font.color.rgb = RGBColor(110, 110, 110)

    doc.add_paragraph()

    prepared = doc.add_paragraph("Prepared for: Tom Short")
    prepared.alignment = WD_ALIGN_PARAGRAPH.CENTER
    prepared.runs[0].font.size = Pt(13)

    doc.add_page_break()

    # Executive Summary
    doc.add_heading("Executive Summary", level=1)

    # Short intro paragraph
    doc.add_paragraph(
        "This executive summary provides a high-level snapshot of your portfolio using live market data, "
        "including valuation, recent performance, diversification, rebalancing needs, and positioning versus key benchmarks."
    )

    # -------------------- Portfolio Snapshot (Table) --------------------
    h_ps = doc.add_heading("Portfolio Snapshot", level=2)
    h_ps.paragraph_format.space_before = Pt(6)

    snapshot_rows = []
    snapshot_rows.append(["Total Value", fmt_dollar(summary_total_value)])

    if summary_target_value is not None:
        snapshot_rows.append(["Target Portfolio Value", fmt_dollar(summary_target_value)])

    snapshot_rows.append(
        ["1D Return", fmt_pct(summary_1d_return) if not np.isnan(summary_1d_return) else "-"]
    )
    snapshot_rows.append(["Total 1D P/L", fmt_dollar(total_1d_pl)])

    snapshot_rows.append(
        ["MTD Return", fmt_pct(port_mtd) if not np.isnan(port_mtd) else "-"]
    )
    snapshot_rows.append(["Total MTD P/L", fmt_dollar(total_mtd_pl)])

    snapshot_rows.append(
        ["YTD Return", fmt_pct(port_ytd) if not np.isnan(port_ytd) else "-"]
    )
    snapshot_rows.append(["Total YTD P/L", fmt_dollar(total_ytd_pl)])

    snapshot_rows.append(["Number of Holdings", str(summary_num_holdings)])

    add_table(
        ["Metric", "Value"],
        snapshot_rows,
        right_align_cols=[1],
    )

    # ----------------- Performance Highlights (Table) -------------------
    h_ph = doc.add_heading("Performance Highlights", level=2)
    h_ph.paragraph_format.space_before = Pt(10)

    perf_rows = [
        ["Top 1M Performer", top_1m_line],
        ["Bottom 1M Performer", bottom_1m_line],
        ["Best 1D Performer", best_1d_line if best_1d_line else "N/A"],
    ]

    add_table(
        ["Metric", "Value"],
        perf_rows,
        right_align_cols=[1],
    )

    # --------------- Risk & Diversification (Table) ---------------------
    h_rd = doc.add_heading("Risk & Diversification", level=2)
    h_rd.paragraph_format.space_before = Pt(10)

    top3_str = fmt_pct_level(top3_pct) if not np.isnan(top3_pct) else "-"
    largest_ac_str = (
        f"{largest_ac_name} ({largest_ac_pct:.2f}%)" if not np.isnan(largest_ac_pct) else "N/A"
    )

    risk_rows = [
        ["Top 3 holdings % of portfolio", top3_str],
        ["Largest asset class", largest_ac_str],
        ["Largest overweight", largest_overweight_str],
        ["Largest underweight", largest_underweight_str],
    ]

    add_table(
        ["Metric", "Value"],
        risk_rows,
        right_align_cols=[1],
    )

    # --------- Rebalancing & Benchmark Positioning (Table) --------------
    h_rb = doc.add_heading("Rebalancing & Benchmark Positioning", level=2)
    h_rb.paragraph_format.space_before = Pt(10)

    pct_target_str = f"{pct_tickers_on_target:.1f}%" if pct_tickers_on_target is not None else "N/A"

    reb_bench_rows = [
        ["Rebalancing need", fmt_dollar(rebalance_need)],
        ["% of tickers on target (±5% band)", pct_target_str],
        ["Vs S&P 500 MTD", vs_sp500_mtd_str],
        ["Vs S&P 500 YTD", vs_sp500_ytd_str],
    ]

    add_table(
        ["Metric", "Value"],
        reb_bench_rows,
        right_align_cols=[1],
    )

    doc.add_page_break()

    # Holdings by Ticker
    doc.add_heading("Portfolio Composition & Strategy", level=1)
    doc.add_heading("Holdings by Ticker", level=2)
    ticker_rows = []
    for _, row in df.iterrows():
        ticker_rows.append(
            [
                row["ticker"],
                row["asset_class"],
                f"{row['shares']:.4f}",
                f"{row['price']:.2f}",
                f"{row['value']:.2f}",
                f"{row['allocation_pct']:.2f}%",
                f"{(row.get('target_pct') or 0.0):.2f}%",
                "-"
                if pd.isna(row.get("contribute_to_target"))
                else fmt_dollar(row["contribute_to_target"]),
            ]
        )

    total_value_sum = df["value"].sum()
    total_contrib_sum = df["contribute_to_target"].fillna(0).sum()

    ticker_rows.append(
        [
            "TOTAL",
            "",
            "",
            "",
            f"{total_value_sum:,.2f}",
            "",
            "",
            fmt_dollar(total_contrib_sum),
        ]
    )

    add_table(
        [
            "Ticker",
            "Asset Class",
            "Shares",
            "Price ($)",
            "Value ($)",
            "Allocation %",
            "Target %",
            "Contribute to Target ($)",
        ],
        ticker_rows,
        right_align_cols=[2, 3, 4, 5, 6, 7],
    )

    # --------- Illustrative Monthly Contribution Schedule (Table) ---------
    if monthly_contrib > 0 and not df.empty:
        need_series = df["contribute_to_target"].clip(lower=0).fillna(0)
        total_need_schedule = float(need_series.sum())

        if total_need_schedule > 0:
            sched_rows = []
            for _, row in df.iterrows():
                gap = float(row.get("contribute_to_target") or 0.0)
                if gap <= 0:
                    continue

                share = gap / total_need_schedule
                monthly_dollars = monthly_contrib * share

                sched_rows.append(
                    [
                        row["ticker"],
                        row["asset_class"],
                        fmt_dollar(gap),
                        fmt_dollar(monthly_dollars),
                        f"{share * 100:.1f}%",
                    ]
                )

            if sched_rows:
                portfolio_months_to_target = total_need_schedule / monthly_contrib

                h_sched = doc.add_heading(
                    "Illustrative Monthly Contribution Schedule", level=2
                )
                h_sched.paragraph_format.space_before = Pt(10)

                add_table(
                    ["Ticker", "Asset Class", "Gap to Target ($)", "Suggested Monthly ($)", "Share of Monthly %"],
                    sched_rows,
                    right_align_cols=[2, 3, 4],
                )

                p_sched_note = doc.add_paragraph(
                    f"At approximately ${monthly_contrib:,.0f}/month, this schedule would allocate "
                    f"contributions proportionally to each holding's shortfall. In total, it would "
                    f"take about {portfolio_months_to_target:.1f} months of contributions to fully "
                    "close the current gaps, assuming markets are flat and this allocation is followed."
                )
                p_sched_note.paragraph_format.space_before = Pt(4)

    # Asset Class Allocation Overview
    h_ac = doc.add_heading("Asset Class Allocation Overview", level=2)
    h_ac.paragraph_format.space_before = Pt(12)

    if asset_targets_df is not None and not asset_targets_df.empty:
        ac_compare = asset_df.merge(
            asset_targets_df, on="asset_class", how="left", suffixes=("", "_target")
        )
        ac_compare["target_pct"] = ac_compare["target_pct"].fillna(0.0)
        ac_compare["delta_pct"] = ac_compare["allocation_pct"] - ac_compare["target_pct"]

        ac_compare = ac_compare.reindex(
            ac_compare["delta_pct"].abs().sort_values(ascending=False).index
        )

        ac_rows = []
        for _, row in ac_compare.iterrows():
            delta_str = fmt_pct(row["delta_pct"])
            ac_rows.append(
                [
                    row["asset_class"],
                    f"{row['value']:.2f}",
                    f"{row['allocation_pct']:.2f}%",
                    f"{row['target_pct']:.2f}%",
                    delta_str,
                ]
            )

        total_ac_value = asset_df["value"].sum()

        ac_rows.append(
            [
                "TOTAL",
                f"{total_ac_value:,.2f}",
                "100.00%",
                f"{asset_targets_df['target_pct'].sum():.2f}%",
                "",
            ]
        )

        add_table(
            ["Asset Class", "Value ($)", "Actual %", "Target %", "Delta %"],
            ac_rows,
            right_align_cols=[1, 2, 3, 4],
        )

    else:
        ac_rows = []
        for _, row in asset_df.iterrows():
            ac_rows.append(
                [
                    row["asset_class"],
                    f"{row['value']:.2f}",
                    f"{row['allocation_pct']:.2f}%",
                ]
            )
        add_table(
            ["Asset Class", "Value ($)", "Allocation %"],
            ac_rows,
            right_align_cols=[1, 2],
        )

    # Next Steps & Ongoing Strategy
    h_ns = doc.add_heading("Next Steps & Ongoing Strategy", level=2)
    h_ns.paragraph_format.space_before = Pt(12)

    add_table(
        ["Focus Area", "Guidance"],
        [
            ["Rebalancing", "Review annually; adjust if allocations drift ±5%."],
            [
                "Contributions",
                f"Automate around ${monthly_contrib:,.0f}/mo into your target allocation.",
            ],
            [
                "Risk Management",
                "Stay within your risk tolerance; increase bonds as you approach major goals.",
            ],
            [
                "Monitoring",
                "Avoid overreacting to short-term volatility; review quarterly or annually.",
            ],
        ],
    )

    # Visuals – Allocation & Diversification
    doc.add_page_break()
    doc.add_heading("Visual Report – Allocation & Diversification", level=1)

    doc.add_heading("Ticker-Level Allocation Breakdown", level=2)
    doc.add_picture(ticker_pie_stream, width=Inches(5.5))
    p = doc.add_paragraph("Figure 1: Allocation by ticker.")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph()
    doc.add_heading("Asset Class Allocation Breakdown", level=2)
    doc.add_picture(asset_pie_stream, width=Inches(5.5))
    p = doc.add_paragraph("Figure 2: Allocation by asset class.")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph()
    doc.add_heading("Sector Allocation Heatmap", level=2)
    doc.add_picture(sector_stream, width=Inches(5.5))
    p = doc.add_paragraph(
        "Figure 3: Approximate sector exposure based on underlying holdings."
    )
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Performance & Growth
    doc.add_page_break()
    doc.add_heading("Performance & Growth", level=1)
    doc.add_heading("Performance vs Benchmarks (MTD & YTD)", level=2)

    bench_df["MTD %"] = pd.to_numeric(bench_df["MTD %"], errors="coerce")
    bench_df["YTD %"] = pd.to_numeric(bench_df["YTD %"], errors="coerce")
    if "Portfolio (Live)" in bench_df["Benchmark"].values:
        port_row = bench_df.loc[bench_df["Benchmark"] == "Portfolio (Live)"].iloc[0]
    else:
        port_row = pd.Series({"MTD %": np.nan, "YTD %": np.nan})

    mtd_rows, ytd_rows = [], []
    for _, row in bench_df.iterrows():
        if row["Benchmark"] == "Portfolio (Live)":
            continue

        port_mtd_val = port_row.get("MTD %")
        bench_mtd_val = row["MTD %"]
        if pd.isna(port_mtd_val) or pd.isna(bench_mtd_val):
            excess_mtd = None
        else:
            excess_mtd = float(port_mtd_val) - float(bench_mtd_val)

        port_ytd_val = port_row.get("YTD %")
        bench_ytd_val = row["YTD %"]
        if pd.isna(port_ytd_val) or pd.isna(bench_ytd_val):
            excess_ytd = None
        else:
            excess_ytd = float(port_ytd_val) - float(bench_ytd_val)

        mtd_rows.append(
            [
                row["Benchmark"],
                fmt_pct(port_mtd_val),
                fmt_pct(bench_mtd_val),
                fmt_pct(excess_mtd),
            ]
        )
        ytd_rows.append(
            [
                row["Benchmark"],
                fmt_pct(port_ytd_val),
                fmt_pct(bench_ytd_val),
                fmt_pct(excess_ytd),
            ]
        )

    # --- MTD Table ---
    p_mtd = doc.add_paragraph("Month-to-date (MTD) comparison:")
    p_mtd.paragraph_format.space_before = Pt(6)

    add_table(
        ["Benchmark", "Portfolio MTD %", "Benchmark MTD %", "Excess MTD %"],
        mtd_rows,
        right_align_cols=[1, 2, 3],
    )

    # --- YTD Table ---
    p_ytd = doc.add_paragraph("Year-to-date (YTD) comparison:")
    p_ytd.paragraph_format.space_before = Pt(12)

    add_table(
        ["Benchmark", "Portfolio YTD %", "Benchmark YTD %", "Excess YTD %"],
        ytd_rows,
        right_align_cols=[1, 2, 3],
    )

    doc.add_paragraph()
    doc.add_heading("Holdings Multi-Horizon Returns", level=2)
    doc.add_paragraph(
        "Performance of each holding over multiple trailing lookback periods "
        "(1D / 1W / 1M / 3M / 6M), ranked by 6-month return. "
        "Note: MTD and YTD performance are shown separately at the portfolio level."
    )

    returns_sorted = returns_df.sort_values("6M %", ascending=False, na_position="last")
    rows = []
    for _, r in returns_sorted.iterrows():
        rows.append(
            [
                r["Ticker"],
                fmt_pct(r["1D %"]),
                fmt_pct(r["1W %"]),
                fmt_pct(r["1M %"]),
                fmt_pct(r["3M %"]),
                fmt_pct(r["6M %"]),
            ]
        )

    # ---- Composite portfolio row (portfolio-level for 1D, weighted for others) ----
    alloc_map = df.set_index("ticker")["allocation_pct"]
    total_row = ["TOTAL"]

    for col in ["1D %", "1W %", "1M %", "3M %", "6M %"]:
        if col == "1D %":
            total_ret = summary_1d_return
        else:
            vals, wts = [], []
            for _, r in returns_sorted.iterrows():
                v = r[col]
                if pd.isna(v):
                    continue
                t = r["Ticker"]
                w = float(alloc_map.get(t, np.nan))
                if np.isnan(w):
                    continue
                vals.append(v)
                wts.append(w)

            total_ret = weighted_avg(vals, wts)

        total_row.append(fmt_pct(total_ret) if not np.isnan(total_ret) else "-")

    rows.append(total_row)

    add_table(
        ["Ticker", "1D %", "1W %", "1M %", "3M %", "6M %"],
        rows,
        right_align_cols=[1, 2, 3, 4, 5],
    )

    sp = doc.add_paragraph()
    sp.paragraph_format.space_before = Pt(0)
    sp.paragraph_format.space_after = Pt(0)

    # Dollar Profit/Loss Table
    dollar_pl_sorted = (
        dollar_pl_df.set_index("Ticker").reindex(returns_sorted["Ticker"]).reset_index()
    )

    p_title = doc.add_heading("Holdings Multi-Horizon Profit/Loss ($)", level=2)
    p_title.paragraph_format.space_before = Pt(2)
    p_title.paragraph_format.keep_with_next = True

    p_desc = doc.add_paragraph(
        "Approximate dollar profit/loss for each holding over the same lookback windows, "
        "using the same return calculations as the percentage table above."
    )
    p_desc.paragraph_format.space_before = Pt(0)
    p_desc.paragraph_format.keep_with_next = True

    pl_rows = []
    for _, r in dollar_pl_sorted.iterrows():
        pl_rows.append(
            [
                r["Ticker"],
                fmt_dollar(r["1D $"]),
                fmt_dollar(r["1W $"]),
                fmt_dollar(r["1M $"]),
                fmt_dollar(r["3M $"]),
                fmt_dollar(r["6M $"]),
            ]
        )

        total_pl_row = ["TOTAL"]

        # 1D $: sum of ticker-level 1D $ so TOTAL matches the rows above
        sum_1d = float(dollar_pl_sorted["1D $"].sum(skipna=True))
        total_pl_row.append(fmt_dollar(sum_1d))

        # Other horizons: sum of per-ticker dollar P/L (unchanged)
        for col in ["1W $", "1M $", "3M $", "6M $"]:
            total_val = float(dollar_pl_sorted[col].sum(skipna=True))
            total_pl_row.append(fmt_dollar(total_val))

        pl_rows.append(total_pl_row)


    add_table(
        ["Ticker", "1D $", "1W $", "1M $", "3M $", "6M $"],
        pl_rows,
        right_align_cols=[1, 2, 3, 4, 5],
    )

    doc.add_paragraph()
    doc.add_heading("Compound Value Breakdown", level=2)
    doc.add_picture(compound_stream, width=Inches(6))

    p = doc.add_paragraph(
        "Figure 4: Illustrates how monthly contributions and compound market growth each drive total portfolio value over 20 years at a 7% annual return."
    )
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(4)

    doc.add_page_break()

    doc.add_heading("20-Year Projection Scenarios", level=2)
    contrib_label = f"+${int(monthly_contrib):,}/mo"
    proj_headers = [
        "Year",
        "5% (Lump)",
        "7% (Lump)",
        "9% (Lump)",
        f"5% ({contrib_label})",
        f"7% ({contrib_label})",
        f"9% ({contrib_label})",
    ]

    # proj_rows were already computed in main and used to build growth chart; we recompute for table
    # to keep behavior identical you'd pass them, but re-building from returns_df isn't straightforward.
    # Easiest and identical to original: compute here too.
    # However to avoid logic drift, assume you pass `proj_rows` later if you want table identical.
    # For now, we just re-use total_value, monthly_contrib and same fixed rates/years.
    rates = [0.05, 0.07, 0.09]
    years = [1, 5, 10, 15, 20]

    def future_value(principal, rate, years_):
        return principal * ((1 + rate / 12) ** (years_ * 12))

    def future_value_with_contrib(principal, rate, years_, monthly_):
        months = years_ * 12
        monthly_rate = rate / 12
        fv_principal = principal * ((1 + monthly_rate) ** months)
        fv_contrib = monthly_ * (((1 + monthly_rate) ** months - 1) / monthly_rate)
        return fv_principal + fv_contrib

    proj_rows = []
    for y in years:
        row = [y]
        for r in rates:
            row.append(round(future_value(total_value, r, y)))
        for r in rates:
            row.append(round(future_value_with_contrib(total_value, r, y, monthly_contrib)))
        proj_rows.append(row)

    proj_rows_for_table = []
    for row in proj_rows:
        formatted_row = [str(row[0])]
        for i in range(1, len(proj_headers)):
            formatted_row.append(f"{row[i]:,.0f}")
        proj_rows_for_table.append(formatted_row)

    add_table(
        proj_headers,
        proj_rows_for_table,
        right_align_cols=list(range(1, len(proj_headers))),
    )

    doc.add_paragraph()
    doc.add_picture(growth_stream, width=Inches(6))
    p = doc.add_paragraph(
        "Figure 5: Variable long-term projections."
    )
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Risk & Volatility  (no page break here anymore)
    doc.add_heading("Risk & Volatility Analysis", level=1)

    h_vol = doc.add_heading("Expected Volatility by Asset Class", level=2)
    h_vol.paragraph_format.space_before = Pt(4)
    h_vol.paragraph_format.space_after = Pt(2)

    doc.add_picture(vol_stream, width=Inches(5.25), height=Inches(3.4))
    p = doc.add_paragraph(
        "Figure 6: Approximate volatility (standard deviation) by asset class."
    )
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(1)
    p.paragraph_format.space_after = Pt(12)

    h_risk = doc.add_heading("Risk vs Expected Return", level=2)
    h_risk.paragraph_format.space_before = Pt(4)
    h_risk.paragraph_format.space_after = Pt(2)

    doc.add_picture(risk_stream, width=Inches(5.25), height=Inches(3.4))
    p = doc.add_paragraph(
        "Figure 7: Trade-off between expected return and volatility by asset class."
    )
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(1)
    p.paragraph_format.space_after = Pt(0)

    # ------------------------- SAVE DOCX + PDF ------------------------

    today_str = get_eastern_now().strftime("%Y-%m-%d")
    docx_name = f"Investment_Report_{today_str}.docx"
    pdf_name = f"Investment_Report_{today_str}.pdf"

    doc.save(docx_name)
    print(f"Report generated: {docx_name}")

    if os.path.exists(pdf_name):
        os.remove(pdf_name)

    if os.path.exists(docx_name):
        os.utime(docx_name, None)

    pdf_created = False

    # Only generate PDF on Windows / macOS; skip on Linux (Colab)
    if sys.platform.startswith("win") or sys.platform == "darwin":
        try:
            from docx2pdf import convert  # lazy import so Colab / Linux don't care
            print("Attempting PDF export via docx2pdf (Word)...")
            convert(docx_name, pdf_name)
            if os.path.exists(pdf_name):
                pdf_created = True
                print(f"✔ PDF created via docx2pdf: {pdf_name}")
            else:
                print("⚠ docx2pdf ran, but PDF file not found.")
        except Exception as e:
            print(f"❌ PDF export failed via docx2pdf: {e}")
            print("Make sure Microsoft Word is installed and docx2pdf is configured.")
    else:
        print("Skipping PDF generation on this platform (DOCX only).")

    if not pdf_created:
        print("⚠ No PDF generated. You still have the DOCX file.")

    # Copy outputs into Drive / other folders
    try:
        dest_dirs = list(EXTRA_OUTPUT_DIRS) if "EXTRA_OUTPUT_DIRS" in globals() else []

        colab_drive_outputs = "/content/drive/MyDrive/Investment Report Outputs"
        if os.path.isdir(colab_drive_outputs) and colab_drive_outputs not in dest_dirs:
            dest_dirs.append(colab_drive_outputs)

        files_to_copy = [docx_name]
        if pdf_created and os.path.exists(pdf_name):
            files_to_copy.append(pdf_name)

        for out_dir in dest_dirs:
            if not os.path.isdir(out_dir):
                continue
            for fname in files_to_copy:
                src = os.path.abspath(fname)
                dst = os.path.join(out_dir, os.path.basename(fname))
                try:
                    shutil.copy2(src, dst)
                    print(f"Copied {fname} -> {dst}")
                except Exception as copy_err:
                    print(f"Could not copy {fname} to {out_dir}: {copy_err}")

        if dest_dirs:
            if pdf_created:
                print("✔ Copied DOCX and PDF to output folders.")
            else:
                print("✔ Copied DOCX to output folders (no PDF).")
    except Exception as outer_err:
        print(f"Post-processing copy step failed: {outer_err}")
