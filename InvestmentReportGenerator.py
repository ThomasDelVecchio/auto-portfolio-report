# ==============================================================
# Live portfolio report with allocation, diversification,
# and real-time MTD / YTD benchmark comparisons
# ==============================================================

import os
import sys
import subprocess
import shutil
from io import BytesIO
from datetime import datetime

import yfinance as yf
import pandas as pd
import numpy as np
import matplotlib.pyplot as plt
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# -------------------------- Timezone Helper --------------------------

try:
    # Python 3.9+ standard library
    from zoneinfo import ZoneInfo

    def get_eastern_now():
        """Return current time in America/New_York (Eastern Time)."""
        return datetime.now(ZoneInfo("America/New_York"))

except Exception:
    # Fallback if zoneinfo isn't available; use pytz if present
    try:
        import pytz
        ET = pytz.timezone("US/Eastern")

        def get_eastern_now():
            """Return current time in US/Eastern using pytz."""
            return datetime.now(ET)

    except Exception:
        # Last-resort fallback: system local time
        def get_eastern_now():
            """Fallback to system local time if timezone libraries not available."""
            return datetime.now()


# -------------------------- Formatting --------------------------

def fmt_pct(x):
    import math
    if x is None or (isinstance(x, float) and math.isnan(x)):
        return "-"
    return f"{x:+.2f}%"


def fmt_dollar(x):
    import math
    if x is None or (isinstance(x, float) and math.isnan(x)):
        return "-"
    return f"${x:,.2f}"


# ----------------------------- CONFIG -----------------------------

def _detect_base_input_dir():
    """
    Universal base directory detection:
      - If PORTFOLIO_INPUT_DIR env var is set, use that
      - Else if running in Colab and Drive folder exists, use:
            /content/drive/MyDrive/Investment Report Inputs
      - Else default to current directory
    """
    env_dir = os.environ.get("PORTFOLIO_INPUT_DIR")
    if env_dir:
        return env_dir

    # Try to detect Colab
    try:
        import google.colab  # type: ignore
        in_colab = True
    except Exception:
        in_colab = False

    if in_colab:
        drive_path = "/content/drive/MyDrive/Investment Report Inputs"
        if os.path.isdir(drive_path):
            return drive_path
        return "/content"

    return "."


BASE_INPUT_DIR = _detect_base_input_dir()

HOLDINGS_CSV = os.path.join(BASE_INPUT_DIR, "sample holdings.csv")
ASSET_TARGETS_CSV = os.path.join(BASE_INPUT_DIR, "targets_asset.csv")  # optional

# How to split an asset-class target across its tickers:
#   "value" -> proportional to current market value (default)
#   "equal" -> equal weight among tickers in that class
TARGET_SPLIT_METHOD = "value"

RISK_FREE_RATE = 0.04
monthly_contrib = 250.0
COLOR_MAIN = ["#2563EB", "#10B981", "#F59E0B", "#6366F1", "#14B8A6"]

# Optional: target total portfolio value (set to a number, or leave as None)
TARGET_PORTFOLIO_VALUE = 50000.0

# Optional: extra folders to copy finished reports into (for local runs)
EXTRA_OUTPUT_DIRS = [r"G:\My Drive\Investment Report Outputs"]

# Illustrative long-run assumptions for Risk/Return views
RISK_RETURN = {
    "US Equities":            {"return": 8.0,  "vol": 15.0},
    "International Equity":   {"return": 8.5,  "vol": 17.0},
    "Emerging Markets":       {"return": 9.0,  "vol": 20.0},
    "Fixed Income":           {"return": 4.0,  "vol": 5.0},
    "Real Estate":            {"return": 6.0,  "vol": 12.0},
    "Energy":                 {"return": 6.5,  "vol": 18.0},
    "Innovation/Tech":        {"return": 10.0, "vol": 25.0},
    "Commodities":            {"return": 6.0,  "vol": 10.0},
    "Gold / Precious Metals": {"return": 5.5,  "vol": 12.0},
    "Digital Assets":         {"return": 11.0, "vol": 70.0},
}


# ----------------------- Helpers / Utilities -----------------------

def _norm_col(df: pd.DataFrame, want: str) -> str:
    want = want.strip().lower()
    for c in df.columns:
        if c.strip().lower() == want:
            return c
    raise ValueError(f"Required column '{want}' not found. Present: {list(df.columns)}")


def normalize_allocations(series: pd.Series) -> pd.Series:
    if len(series) == 0:
        return series
    rounded = series.round(2)
    diff = 100.00 - float(rounded.sum())
    rounded.iloc[-1] = round(float(rounded.iloc[-1]) + diff, 2)
    return rounded


def get_live_price(ticker: str) -> float:
    data = yf.Ticker(ticker).history(period="1d")
    if data.empty:
        data = yf.Ticker(ticker).history(period="5d")
        if data.empty:
            raise ValueError(f"No price data for {ticker}")
    return float(data["Close"].iloc[-1])


def get_return_pct(ticker, start_date, end_date):
    try:
        data = yf.download(ticker, start=start_date, end=end_date, progress=False)
        if data.empty:
            data = yf.download(ticker, period="60d", progress=False)
        if data.empty:
            return np.nan
        close = data["Close"]
        end = float(close.iloc[-1])
        start = float(close.iloc[0])
        return (end / start - 1.0) * 100.0
    except Exception:
        return np.nan


def get_1d_return_pct(ticker):
    """
    1D return = (current price / prior day's close) - 1
    Uses period='2d' so we always get yesterday close + today's live price.
    """
    try:
        data = yf.download(ticker, period="2d", progress=False)
        if data.empty or len(data) < 2:
            return np.nan
        close = data["Close"].astype(float)
        prev_close = close.iloc[-2]
        last_price = close.iloc[-1]
        return (last_price / prev_close - 1.0) * 100.0
    except Exception:
        return np.nan


def read_asset_targets(path: str) -> pd.DataFrame | None:
    try:
        df = pd.read_csv(path)
    except Exception:
        return None
    try:
        ac = _norm_col(df, "asset_class")
        tp = _norm_col(df, "target_pct")
    except Exception:
        return None
    out = df.rename(columns={ac: "asset_class", tp: "target_pct"}).copy()
    out["asset_class"] = out["asset_class"].astype(str).str.strip()
    out["target_pct"] = pd.to_numeric(out["target_pct"], errors="coerce").fillna(0.0)
    if out["target_pct"].sum() > 0:
        out["target_pct"] = out["target_pct"] / out["target_pct"].sum() * 100.0
        out["target_pct"] = normalize_allocations(out["target_pct"])
    return out


def build_ticker_targets(df_holdings: pd.DataFrame,
                         df_asset_targets: pd.DataFrame | None,
                         split_method: str = "value") -> dict:
    """
    Priority:
      1) If holdings has per-ticker 'target_pct', use those (normalized).
      2) Else if df_asset_targets provided, distribute to tickers within each asset_class.
      3) Else equal-weight all tickers.
    """
    cols_l = [c.lower().strip() for c in df_holdings.columns]
    tickers = df_holdings["ticker"]

    # 1) Per-ticker targets in holdings
    if "target_pct" in cols_l:
        tcol = df_holdings.columns[cols_l.index("target_pct")]
        tgt = pd.to_numeric(df_holdings[tcol], errors="coerce").fillna(0.0)
        if tgt.sum() > 0:
            scaled = tgt / tgt.sum() * 100.0
            scaled = normalize_allocations(scaled.reset_index(drop=True))
            return dict(zip(tickers, scaled.tolist()))

    # 2) Asset-class targets
    if df_asset_targets is not None and not df_asset_targets.empty:
        merged = df_holdings.merge(
            df_asset_targets, on="asset_class", how="left", suffixes=("", "_ac")
        )
        targets = {}
        for ac, chunk in merged.groupby("asset_class", dropna=False):
            ac_target = float(chunk["target_pct"].iloc[0]) if "target_pct" in chunk else 0.0
            if ac_target <= 0 or len(chunk) == 0:
                continue
            if split_method == "equal":
                each = ac_target / len(chunk)
                for _, r in chunk.iterrows():
                    targets[r["ticker"]] = targets.get(r["ticker"], 0.0) + each
            else:  # proportional to current value
                vals = chunk["value"].clip(lower=0.0)
                denom = float(vals.sum())
                if denom == 0:
                    each = ac_target / len(chunk)
                    for _, r in chunk.iterrows():
                        targets[r["ticker"]] = targets.get(r["ticker"], 0.0) + each
                else:
                    for _, r in chunk.iterrows():
                        w = float(r["value"]) / denom if denom > 0 else 1.0 / len(chunk)
                        targets[r["ticker"]] = targets.get(r["ticker"], 0.0) + (ac_target * w)
        if targets:
            s = sum(targets.values())
            if s > 0:
                for k in list(targets.keys()):
                    targets[k] = targets[k] / s * 100.0
            keys = list(targets.keys())
            vals = [round(v, 2) for v in targets.values()]
            diff = round(100.0 - sum(vals), 2)
            vals[-1] = round(vals[-1] + diff, 2)
            return dict(zip(keys, vals))

    # 3) Equal-weight
    n = len(df_holdings)
    if n == 0:
        return {}
    eq = [round(100.0 / n, 2)] * n
    eq[-1] = round(eq[-1] + (100.0 - sum(eq)), 2)
    return dict(zip(tickers, eq))


# --------------------- 1) LOAD HOLDINGS + PRICES ---------------------

raw = pd.read_csv(HOLDINGS_CSV)

tcol = _norm_col(raw, "ticker")
scol = _norm_col(raw, "shares")

df = raw.rename(columns={tcol: "ticker", scol: "shares"}).copy()
df["ticker"] = df["ticker"].astype(str).str.upper().str.strip()
df["shares"] = pd.to_numeric(df["shares"], errors="coerce").fillna(0.0)

# asset_class optional
if "asset_class" not in [c.strip().lower() for c in df.columns]:
    df["asset_class"] = "Unknown"
else:
    ac_col = [c for c in df.columns if c.strip().lower() == "asset_class"][0]
    df.rename(columns={ac_col: "asset_class"}, inplace=True)

# Fetch prices; drop tickers that fail
prices = []
for t in df["ticker"]:
    try:
        prices.append(get_live_price(t))
    except Exception:
        prices.append(np.nan)

df["price"] = prices
df = df.dropna(subset=["price"]).reset_index(drop=True)

df["value"] = df["shares"] * df["price"]
total_value = float(df["value"].sum()) if len(df) else 0.0
df["allocation_pct"] = np.where(total_value > 0, (df["value"] / total_value * 100.0), 0.0)
df["allocation_pct"] = normalize_allocations(df["allocation_pct"])

# ------------- 2) ASSET-CLASS SUMMARY (entire portfolio) -------------

asset_df = df.groupby("asset_class", as_index=False)["value"].sum()
asset_df["allocation_pct"] = np.where(
    total_value > 0, (asset_df["value"] / total_value * 100.0), 0.0
)
asset_df["allocation_pct"] = normalize_allocations(asset_df["allocation_pct"])

# ------------------ TARGETS (user-determined) -------------------

asset_targets_df = read_asset_targets(ASSET_TARGETS_CSV)

TICKER_TARGETS_PCT = build_ticker_targets(df.copy(), asset_targets_df, TARGET_SPLIT_METHOD)

core_df = df.copy()
core_total = TARGET_PORTFOLIO_VALUE if TARGET_PORTFOLIO_VALUE is not None else total_value

core_df["core_allocation_pct"] = np.where(
    total_value > 0,
    (core_df["value"] / total_value * 100.0),
    0.0,
)

core_df["target_pct"] = (
    core_df["ticker"]
    .map(TICKER_TARGETS_PCT)
    .astype(float)
    .fillna(0.0)
)

core_df["target_value"] = core_total * (core_df["target_pct"] / 100.0)

for col in ["target_pct", "target_value", "core_allocation_pct"]:
    if col in df.columns:
        df.drop(columns=[col], inplace=True)

df = df.merge(
    core_df[["ticker", "target_value", "target_pct", "core_allocation_pct"]],
    on="ticker",
    how="left",
)

df["delta_to_target_raw"] = df["target_value"] - df["value"]

df["contribute_to_target"] = np.where(
    df["delta_to_target_raw"].isna(),
    np.nan,
    np.where(df["delta_to_target_raw"] > 0, df["delta_to_target_raw"], 0.0),
)


# ==============================================================
# 3) REAL ETF-BASED SECTOR WEIGHTS (ETF map + Yahoo for singles)
# ==============================================================

ETF_SECTOR_MAP = {
    "VOO": {
        "Tech": 29.0,
        "Financials": 13.0,
        "Health Care": 13.0,
        "Industrials": 8.0,
        "Consumer Disc.": 10.0,
        "Comm Services": 9.0,
        "Energy": 4.0,
        "Materials": 2.5,
        "Real Estate": 2.5,
        "Utilities": 2.5,
    },
    "QQQ": {
        "Tech": 55.0,
        "Consumer Disc.": 17.0,
        "Comm Services": 15.0,
        "Health Care": 7.0,
        "Industrials": 3.0,
        "Other": 3.0,
    },
    "VXUS": {
        "Financials": 20.0,
        "Industrials": 15.0,
        "Consumer Disc.": 12.0,
        "Tech": 11.0,
        "Health Care": 9.0,
        "Materials": 8.0,
        "Energy": 6.0,
        "Real Estate": 6.0,
        "Utilities": 4.0,
        "Comm Services": 4.0,
    },
    "BND": {"Fixed Income": 100.0},
    "GLD": {"Precious Metals": 100.0},
    "FBTC": {"Digital Assets": 100.0},
}

# Normalize Yahoo sector names to match ETF naming convention
SECTOR_NAME_NORMALIZE = {
    "Technology": "Tech",
    "Information Technology": "Tech",
    "Financial Services": "Financials",
    "Consumer Cyclical": "Consumer Disc.",
    "Communication Services": "Comm Services",
}

def normalize_sector_name(name: str) -> str:
    if not isinstance(name, str):
        return "Other"
    base = name.strip()
    return SECTOR_NAME_NORMALIZE.get(base, base)


portfolio_sectors = {}
sector_cache = {}

for _, row in df.iterrows():
    t = row["ticker"]
    w = float(row["allocation_pct"])

    # 1) ETF tickers → use static breakdown
    if t in ETF_SECTOR_MAP:
        for sector, pct in ETF_SECTOR_MAP[t].items():
            portfolio_sectors[sector] = portfolio_sectors.get(sector, 0.0) + (w * pct / 100.0)
        continue

    # 2) Single stocks → use Yahoo Finance
    if t in sector_cache:
        sector = sector_cache[t]
    else:
        try:
            info = yf.Ticker(t).info
            raw_sector = info.get("sector")
        except Exception:
            raw_sector = None

        sector = normalize_sector_name(raw_sector) if raw_sector else "Other"
        sector_cache[t] = sector

    portfolio_sectors[sector] = portfolio_sectors.get(sector, 0.0) + w

# Build DataFrame and normalize
sector_df_static = pd.DataFrame(
    list(portfolio_sectors.items()),
    columns=["Sector", "Weight"]
)

total = sector_df_static["Weight"].sum()
if total > 0:
    sector_df_static["Weight"] = (sector_df_static["Weight"] / total * 100.0).round(2)

sector_df_static = sector_df_static.sort_values("Weight", ascending=False).reset_index(drop=True)




# ---------------------- 4) SECTOR HEATMAP CHART ----------------------

plt.figure(figsize=(6, 5))
plt.barh(
    sector_df_static["Sector"], sector_df_static["Weight"],
    color=plt.cm.Blues(np.linspace(0.4, 0.9, len(sector_df_static)))
)
plt.xlabel("Portfolio Exposure (%)")
plt.title("Sector Allocation Heatmap", fontsize=12, weight="bold")
plt.gca().invert_yaxis()
plt.tight_layout()
sector_stream = BytesIO()
plt.savefig(sector_stream, format="png", bbox_inches="tight", facecolor="white")
sector_stream.seek(0)
plt.close()


# -------------------- 5) BENCHMARK DATA (MTD & YTD) -------------------

benchmarks = {
    "S&P 500": "^GSPC",
    "Global 60/40": "AOR",
    "Conservative 40/60": "AOK",
}

today = get_eastern_now()

# Last business day of the prior month
start_month = (today.replace(day=1) - pd.offsets.BMonthEnd(1))

# Last business day before Jan 1
start_year = (datetime(today.year, 1, 1) - pd.offsets.BMonthEnd(1))


bench_rows = []
for name, ticker in benchmarks.items():
    mtd = get_return_pct(ticker, start_month, today)
    ytd = get_return_pct(ticker, start_year, today)
    bench_rows.append(
        {
            "Benchmark": name,
            "MTD %": round(mtd, 2) if not np.isnan(mtd) else np.nan,
            "YTD %": round(ytd, 2) if not np.isnan(ytd) else np.nan,
        }
    )

port_changes_mtd, port_changes_ytd, weights = [], [], []
for _, row in df.iterrows():
    t = row["ticker"]
    w = row["allocation_pct"]
    mtd = get_return_pct(t, start_month, today)
    ytd = get_return_pct(t, start_year, today)
    if np.isnan(mtd) or np.isnan(ytd):
        continue
    port_changes_mtd.append(mtd)
    port_changes_ytd.append(ytd)
    weights.append(w)


def weighted_avg(values, weights):
    if not values:
        return np.nan
    v = np.array(values, dtype=float)
    w = np.array(weights, dtype=float)

    s = w.sum()
    if s == 0:
        return np.nan

    w = w / s  # normalize
    return float((v * w).sum())



port_mtd = weighted_avg(port_changes_mtd, weights)
port_ytd = weighted_avg(port_changes_ytd, weights)

bench_rows.insert(
    0,
    {
        "Benchmark": "Portfolio (Live)",
        "MTD %": round(port_mtd, 2) if not np.isnan(port_mtd) else np.nan,
        "YTD %": round(port_ytd, 2) if not np.isnan(port_ytd) else np.nan,
    },
)

bench_df = pd.DataFrame(bench_rows)


# --------- 6) HOLDINGS MULTI-HORIZON RETURNS (1D, 1W, 1M, 3M, 6M) ---------

horizons = {
    "1D %": 1,
    "1W %": 7,
    "1M %": 30,
    "3M %": 90,
    "6M %": 180,
}

returns_rows = []
for _, row in df.iterrows():
    t = row["ticker"]
    r = {"Ticker": t}

    for label, days in horizons.items():
        if label == "1D %":
            price_data = yf.Ticker(t).history(period="2d")
            if len(price_data) >= 2:
                prev_close = float(price_data["Close"].iloc[-2])
                current = float(price_data["Close"].iloc[-1])
                r[label] = round(((current / prev_close) - 1) * 100, 2)
            else:
                r[label] = np.nan
        else:
            start = today - pd.Timedelta(days=days)
            val = get_return_pct(t, start, today)
            r[label] = round(val, 2) if not np.isnan(val) else np.nan

    returns_rows.append(r)

returns_df = pd.DataFrame(returns_rows)


# ---------- Dollar Profit/Loss for Each Horizon (using same math) ----------

def dollar_pl_from_return(current_value, pct):
    """
    Convert a percentage return (end/start - 1) into a dollar P/L
    using the current value as 'end'.

    start_value = current_value / (1 + r)
    P/L = current_value - start_value
    """
    import math

    if pct is None or (isinstance(pct, float) and math.isnan(pct)):
        return np.nan

    r = float(pct) / 100.0
    if r <= -0.9999:
        return -current_value

    try:
        start_val = current_value / (1.0 + r)
    except ZeroDivisionError:
        return np.nan

    return current_value - start_val


returns_by_ticker = returns_df.set_index("Ticker")

dollar_pl_rows = []
for _, h_row in df.iterrows():
    t = h_row["ticker"]
    if t not in returns_by_ticker.index:
        continue

    cur_val = float(h_row["value"])
    r_row = returns_by_ticker.loc[t]

    row = {
        "Ticker": t,
        "1D $": dollar_pl_from_return(cur_val, r_row.get("1D %")),
        "1W $": dollar_pl_from_return(cur_val, r_row.get("1W %")),
        "1M $": dollar_pl_from_return(cur_val, r_row.get("1M %")),
        "3M $": dollar_pl_from_return(cur_val, r_row.get("3M %")),
        "6M $": dollar_pl_from_return(cur_val, r_row.get("6M %")),
    }
    dollar_pl_rows.append(row)

dollar_pl_df = pd.DataFrame(dollar_pl_rows)


# ---------------- 7) LONG-TERM PROJECTIONS (20 YEARS) -----------------

rates = [0.05, 0.07, 0.09]
years = [1, 5, 10, 15, 20]


def future_value(principal, rate, years):
    return principal * ((1 + rate / 12) ** (years * 12))


def future_value_with_contrib(principal, rate, years, monthly):
    months = years * 12
    monthly_rate = rate / 12
    fv_principal = principal * ((1 + monthly_rate) ** months)
    fv_contrib = monthly * (((1 + monthly_rate) ** months - 1) / monthly_rate)
    return fv_principal + fv_contrib


proj_rows = []
for y in years:
    row = [y]
    for r in rates:
        row.append(round(future_value(total_value, r, y)))
    for r in rates:
        row.append(round(future_value_with_contrib(total_value, r, y, monthly_contrib)))
    proj_rows.append(row)

plt.figure(figsize=(6, 4))
for i, r in enumerate(rates):
    plt.plot(
        years,
        [future_value(total_value, r, y) for y in years],
        label=f"{int(r*100)}% Lump Sum",
        linestyle="-",
    )
    plt.plot(
        years,
        [future_value_with_contrib(total_value, r, y, monthly_contrib) for y in years],
        label=f"{int(r*100)}% + ${int(monthly_contrib)}/mo",
        linestyle="--",
    )
plt.title("Portfolio Growth Projections (20-Year Scenarios)", fontsize=12, weight="bold")
plt.xlabel("Years")
plt.ylabel("Portfolio Value ($)")
plt.grid(alpha=0.3)
plt.legend(fontsize=8)
plt.tight_layout()
growth_stream = BytesIO()
plt.savefig(growth_stream, format="png", bbox_inches="tight", facecolor="white")
growth_stream.seek(0)
plt.close()


# -------- 8) COMPOUND VALUE BREAKDOWN (CONTRIB VS GROWTH) -------------

years_compound = list(range(0, 21))
rate_for_compound = 0.07
initial_balance = total_value

contrib_values, growth_values = [], []
for y in years_compound:
    total_with_contrib = future_value_with_contrib(
        initial_balance, rate_for_compound, y, monthly_contrib
    )
    # All dollars you've put in: today's balance + future contributions
    total_contrib = initial_balance + monthly_contrib * 12 * y
    contrib_values.append(total_contrib)
    growth_values.append(max(total_with_contrib - total_contrib, 0))


plt.figure(figsize=(6, 4))
plt.stackplot(
    years_compound,
    contrib_values,
    growth_values,
    labels=["Contributions", "Growth"],
    colors=[COLOR_MAIN[0], COLOR_MAIN[1]],
    alpha=0.85,
)

plt.title(
    "Contributions vs Growth",
    fontsize=12,
    weight="bold",
)

plt.xlabel("Years")
plt.ylabel("Portfolio Value ($)")
plt.legend(loc="upper left", fontsize=8)
plt.tight_layout()
compound_stream = BytesIO()
plt.savefig(compound_stream, format="png", bbox_inches="tight", facecolor="white")
compound_stream.seek(0)
plt.close()


# -------------- 9) PERFORMANCE VS BENCHMARKS (CHART) ------------------

bench_df["MTD %"] = pd.to_numeric(bench_df["MTD %"], errors="coerce")
bench_df["YTD %"] = pd.to_numeric(bench_df["YTD %"], errors="coerce")
bench_plot = bench_df.dropna(subset=["MTD %", "YTD %"], how="all").reset_index(drop=True)

if bench_plot.empty:
    plt.figure(figsize=(6, 4))
    plt.text(
        0.5,
        0.5,
        "Benchmark return data unavailable.\n(Check connection or market days.)",
        ha="center",
        va="center",
        fontsize=9,
    )
    plt.axis("off")
    plt.tight_layout()
else:
    x = np.arange(len(bench_plot))
    width = 0.35
    mtd_vals = bench_plot["MTD %"].to_numpy(dtype=float)
    ytd_vals = bench_plot["YTD %"].to_numpy(dtype=float)

    plt.figure(figsize=(6, 4))
    mtd_bars = plt.bar(x - width / 2, mtd_vals, width, label="MTD", color=COLOR_MAIN[0])
    ytd_bars = plt.bar(x + width / 2, ytd_vals, width, label="YTD", color=COLOR_MAIN[1])
    plt.xticks(x, bench_plot["Benchmark"], rotation=20, ha="right")
    plt.ylabel("Return (%)")
    plt.title("Portfolio vs Benchmarks (MTD & YTD)", fontsize=12, weight="bold")
    plt.legend(fontsize=8)
    plt.grid(axis="y", alpha=0.3)

    vals = np.concatenate(
        [mtd_vals[~np.isnan(mtd_vals)], ytd_vals[~np.isnan(ytd_vals)]]
    ) if (len(mtd_vals) + len(ytd_vals)) else np.array([])
    if vals.size > 0:
        vmin, vmax = float(np.min(vals)), float(np.max(vals))
        if vmin == vmax:
            vmin -= 1.0
            vmax += 1.0
        else:
            pad = (vmax - vmin) * 0.2
            vmin -= pad
            vmax += pad
        plt.ylim(vmin, vmax)

    def label_bars(bars, vals):
        for bar, val in zip(bars, vals):
            if np.isnan(val):
                continue
            height = bar.get_height()
            plt.text(
                bar.get_x() + bar.get_width() / 2,
                height,
                f"{val:+.1f}%",
                ha="center",
                va="bottom",
                fontsize=7,
            )

    label_bars(mtd_bars, mtd_vals)
    label_bars(ytd_bars, ytd_vals)
    plt.tight_layout()


# ------------------ 10) RISK & VOLATILITY (CLEANED) ------------------

present_assets = sorted(set(asset_df["asset_class"]))

risk_rows = []
for ac in present_assets:
    ac_norm = ac.strip().lower()

    # normalize plurals (Equities → Equity)
    ac_norm = ac_norm.replace("equities", "equity")

    # Try direct exact matches first (case-insensitive)
    base = None
    for k, v in RISK_RETURN.items():
        key_norm = k.lower().replace("equities", "equity")
        if ac_norm.startswith(key_norm):
            base = v
            break

    # Fallback pattern matching
    if base is None:
        if "international" in ac_norm:
            base = RISK_RETURN.get("International Equity")
        elif "emerging" in ac_norm:
            base = RISK_RETURN.get("Emerging Markets")
        elif "gold" in ac_norm or "precious" in ac_norm:
            base = RISK_RETURN.get("Gold / Precious Metals")
        elif "fixed" in ac_norm or "bond" in ac_norm:
            base = RISK_RETURN.get("Fixed Income")
        elif "real" in ac_norm:
            base = RISK_RETURN.get("Real Estate")
        elif "energy" in ac_norm:
            base = RISK_RETURN.get("Energy")
        elif "tech" in ac_norm or "innovation" in ac_norm:
            base = RISK_RETURN.get("Innovation/Tech")
        elif "commodit" in ac_norm:
            base = RISK_RETURN.get("Commodities")

    # Only append if we successfully matched something
    if base:
        risk_rows.append({"asset_class": ac, "vol": base["vol"], "ret": base["return"]})

# Output DataFrame
risk_df = pd.DataFrame(risk_rows)


plt.figure(figsize=(6, 4))
plt.bar(risk_df["asset_class"], risk_df["vol"], color=COLOR_MAIN[0])
plt.title("Expected Volatility by Asset Class", fontsize=12, weight="bold")
plt.ylabel("Standard Deviation (%)")
plt.xticks(rotation=25, ha="right")
if len(risk_df) and risk_df["vol"].max() == risk_df["vol"].min():
    ymin = risk_df["vol"].min() - 1
    ymax = risk_df["vol"].max() + 1
else:
    ymin = max(0, float(risk_df["vol"].min()) - 2)
    ymax = float(risk_df["vol"].max()) + 4
plt.ylim(ymin, ymax)
plt.grid(axis="y", alpha=0.25)
plt.tight_layout()
vol_stream = BytesIO()
plt.savefig(vol_stream, format="png", bbox_inches="tight", facecolor="white")
vol_stream.seek(0)
plt.close()

plt.figure(figsize=(6, 4))
plt.scatter(
    risk_df["vol"],
    risk_df["ret"],
    s=80,
    edgecolors="black",
    linewidths=0.6,
    alpha=0.9,
    color=COLOR_MAIN[1],
)
for _, row in risk_df.iterrows():
    plt.annotate(
        row["asset_class"],
        (row["vol"] + 0.4, row["ret"]),
        fontsize=8,
        weight="bold",
    )
plt.title("Risk vs Expected Return by Asset Class", fontsize=12, weight="bold")
plt.xlabel("Volatility (Std Dev %)")
plt.ylabel("Expected Annual Return (%)")
if len(risk_df):
    xpad = (
        (risk_df["vol"].max() - risk_df["vol"].min()) * 0.2
        if risk_df["vol"].max() != risk_df["vol"].min()
        else 3
    )
    ypad = (
        (risk_df["ret"].max() - risk_df["ret"].min()) * 0.2
        if risk_df["ret"].max() != risk_df["ret"].min()
        else 2
    )
    plt.xlim(max(0, risk_df["vol"].min() - xpad), risk_df["vol"].max() + xpad)
    plt.ylim(max(0, risk_df["ret"].min() - ypad), risk_df["ret"].max() + ypad)
plt.grid(alpha=0.3)
plt.tight_layout()
risk_stream = BytesIO()
plt.savefig(risk_stream, format="png", bbox_inches="tight", facecolor="white")
risk_stream.seek(0)
plt.close()


# ---------------------- 11) ALLOCATION PIE CHARTS ----------------------

# --- Ticker-level allocation pie ---
plt.figure(figsize=(6, 6))
plt.pie(
    df["allocation_pct"],
    labels=df["ticker"],
    autopct="%1.2f%%",
    startangle=90,
    pctdistance=0.85,
    labeldistance=1.05,
    textprops={"fontsize": 9},
)
plt.title("Portfolio Allocation by Ticker", fontsize=12, weight="bold")

ax = plt.gca()
ax.set_aspect("equal", adjustable="box")
plt.subplots_adjust(left=0.15, right=0.85, top=0.85, bottom=0.15)

ticker_pie_stream = BytesIO()
plt.savefig(ticker_pie_stream, format="png", bbox_inches="tight", facecolor="white")
ticker_pie_stream.seek(0)
plt.close()

# --- Asset-class allocation pie ---

# Ensure all labels are strings before replacement
ac_labels = []
for name in asset_df["asset_class"]:
    label = str(name)  # <-- FIX: cast to string safely
    label = label.replace("International Equities", "Intl. Equities")
    label = label.replace("Precious Metals", "Precious\nMetals")
    label = label.replace("Fixed Income", "Fixed\nIncome")
    label = label.replace("Global Bonds", "Global\nBonds")
    ac_labels.append(label)

plt.figure(figsize=(6, 6))
plt.pie(
    asset_df["allocation_pct"],
    labels=ac_labels,
    autopct="%1.2f%%",
    startangle=90,
    pctdistance=0.80,         # bring % labels inward
    labeldistance=1.03,       # bring text inward
    textprops={"fontsize": 8} # smaller font
)
plt.title("Asset Class Allocation", fontsize=12, weight="bold")

ax = plt.gca()
ax.set_aspect("equal", adjustable="box")
plt.subplots_adjust(left=0.15, right=0.85, top=0.85, bottom=0.15)

asset_pie_stream = BytesIO()
plt.savefig(asset_pie_stream, format="png", bbox_inches="tight", facecolor="white")
asset_pie_stream.seek(0)
plt.close()


# ------------------------- 12) BUILD WORD DOC --------------------------

# ---- Executive Summary Metrics ----

# Portfolio snapshot
summary_total_value = total_value
summary_target_value = TARGET_PORTFOLIO_VALUE
summary_num_holdings = len(df)

# Total 1M P/L (sum of per-ticker 1M dollars)
total_1m_pl = float(dollar_pl_df["1M $"].sum(skipna=True)) if not dollar_pl_df.empty else float("nan")

# Join returns and $ P/L for ranking
ret_pl = returns_df.merge(dollar_pl_df, on="Ticker", how="left")

top_1m_line = "N/A"
bottom_1m_line = "N/A"
best_1d_line = None  # optional

if not ret_pl.empty:
    # Top & bottom 1M performers
    valid_1m = ret_pl.dropna(subset=["1M %"])
    if not valid_1m.empty:
        top_1m_row = valid_1m.sort_values("1M %", ascending=False).iloc[0]
        top_1m_line = (
            f"{top_1m_row['Ticker']} "
            f"({fmt_pct(top_1m_row['1M %'])}, "
            f"{fmt_dollar(top_1m_row.get('1M $'))})"
        )

        bottom_1m_row = valid_1m.sort_values("1M %", ascending=True).iloc[0]
        bottom_1m_line = (
            f"{bottom_1m_row['Ticker']} "
            f"({fmt_pct(bottom_1m_row['1M %'])}, "
            f"{fmt_dollar(bottom_1m_row.get('1M $'))})"
        )

    # Best 1D performer (optional)
    valid_1d = ret_pl.dropna(subset=["1D %"])
    if not valid_1d.empty:
        best_1d_row = valid_1d.sort_values("1D %", ascending=False).iloc[0]
        best_1d_line = (
            f"{best_1d_row['Ticker']} "
            f"({fmt_pct(best_1d_row['1D %'])}, "
            f"{fmt_dollar(best_1d_row.get('1D $'))})"
        )

# Top 3 holdings concentration
if summary_num_holdings > 0:
    top3_pct = float(df["allocation_pct"].nlargest(min(3, summary_num_holdings)).sum())
else:
    top3_pct = float("nan")

# Largest asset class by actual allocation
if not asset_df.empty:
    largest_ac_row = asset_df.sort_values("allocation_pct", ascending=False).iloc[0]
    largest_ac_name = largest_ac_row["asset_class"]
    largest_ac_pct = float(largest_ac_row["allocation_pct"])
else:
    largest_ac_name = "N/A"
    largest_ac_pct = float("nan")

# Overweight / underweight asset classes vs targets
largest_overweight_str = "N/A"
largest_underweight_str = "N/A"

if asset_targets_df is not None and not asset_targets_df.empty:
    ac_compare_for_summary = asset_df.merge(
        asset_targets_df, on="asset_class", how="left", suffixes=("", "_target")
    )
    ac_compare_for_summary["target_pct"] = ac_compare_for_summary["target_pct"].fillna(0.0)
    ac_compare_for_summary["delta_pct"] = (
        ac_compare_for_summary["allocation_pct"] - ac_compare_for_summary["target_pct"]
    )

    ow = ac_compare_for_summary[ac_compare_for_summary["delta_pct"] > 0]
    uw = ac_compare_for_summary[ac_compare_for_summary["delta_pct"] < 0]

    if not ow.empty:
        ow_row = ow.sort_values("delta_pct", ascending=False).iloc[0]
        largest_overweight_str = (
            f"{ow_row['asset_class']} "
            f"({ow_row['allocation_pct']:.2f}% vs {ow_row['target_pct']:.2f}%)"
        )

    if not uw.empty:
        uw_row = uw.sort_values("delta_pct", ascending=True).iloc[0]
        largest_underweight_str = (
            f"{uw_row['asset_class']} "
            f"({uw_row['allocation_pct']:.2f}% vs {uw_row['target_pct']:.2f}%)"
        )

# Rebalancing need & % of tickers on target (within ±5% band)
rebalance_need = float(df["contribute_to_target"].fillna(0).sum()) if not df.empty else 0.0

pct_tickers_on_target = None
if "target_pct" in df.columns and not df["target_pct"].isna().all():
    has_target = df["target_pct"].fillna(0) > 0
    n_with_target = int(has_target.sum())
    if n_with_target > 0:
        drift = (df.loc[has_target, "allocation_pct"] - df.loc[has_target, "target_pct"]).abs()
        within_band = drift <= 5.0  # within ±5% band
        pct_tickers_on_target = float(within_band.mean() * 100.0)

# Benchmark positioning vs S&P 500 (excess returns MTD/YTD)
vs_sp500_mtd_str = "N/A"
vs_sp500_ytd_str = "N/A"

try:
    bench_df["MTD %"] = pd.to_numeric(bench_df["MTD %"], errors="coerce")
    bench_df["YTD %"] = pd.to_numeric(bench_df["YTD %"], errors="coerce")

    if (
        "Portfolio (Live)" in bench_df["Benchmark"].values
        and "S&P 500" in bench_df["Benchmark"].values
    ):
        port_row = bench_df.loc[bench_df["Benchmark"] == "Portfolio (Live)"].iloc[0]
        sp_row = bench_df.loc[bench_df["Benchmark"] == "S&P 500"].iloc[0]

        if not pd.isna(port_row["MTD %"]) and not pd.isna(sp_row["MTD %"]):
            vs_sp500_mtd = float(port_row["MTD %"] - sp_row["MTD %"])
            vs_sp500_mtd_str = fmt_pct(vs_sp500_mtd)

        if not pd.isna(port_row["YTD %"]) and not pd.isna(sp_row["YTD %"]):
            vs_sp500_ytd = float(port_row["YTD %"] - sp_row["YTD %"])
            vs_sp500_ytd_str = fmt_pct(vs_sp500_ytd)
except Exception:
    pass

# Now create the document
doc = Document()
style = doc.styles["Normal"]
style.font.name = "Calibri"
style._element.rPr.rFonts.set(qn("w:eastAsia"), "Calibri")
style.font.size = Pt(11)


def add_table(headers, rows, right_align_cols=None):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Light Grid Accent 1"

    hdr_cells = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr_cells[i].text = h
        for p in hdr_cells[i].paragraphs:
            for r in p.runs:
                r.bold = True

    for row_data in rows:
        row_cells = table.add_row().cells
        for i, val in enumerate(row_data):
            row_cells[i].text = str(val)

    # Repeat header row on each page
    hdr_row = table.rows[0]
    tr = hdr_row._tr
    trPr = tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    trPr.append(tbl_header)

    # Prevent splitting rows; keep rows together
    row_count = len(table.rows)
    for idx, row in enumerate(table.rows):
        tr = row._tr
        trPr = tr.get_or_add_trPr()

        cant_split = OxmlElement("w:cantSplit")
        trPr.append(cant_split)

        for cell in row.cells:
            for p in cell.paragraphs:
                pf = p.paragraph_format
                pf.keep_together = True
                pf.keep_with_next = (idx < row_count - 1)

    # Right-align numeric columns
    if right_align_cols is not None:
        for row in table.rows[1:]:
            for col_idx in right_align_cols:
                if col_idx < len(row.cells):
                    cell = row.cells[col_idx]
                    for p in cell.paragraphs:
                        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT

    return table


# Cover
doc.add_paragraph("\n\n\n\n\n")

cover = doc.add_paragraph("Comprehensive Investment Report — Live Portfolio")
cover.runs[0].bold = True
cover.runs[0].font.size = Pt(26)
cover.alignment = WD_ALIGN_PARAGRAPH.CENTER

subtitle = doc.add_paragraph(
    "Automated Summary, Diversification, and Performance vs Benchmarks"
)
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
subtitle.runs[0].font.size = Pt(14)
subtitle.runs[0].font.color.rgb = RGBColor(70, 130, 180)

timestamp_str = get_eastern_now().strftime("%A, %B %d, %Y — %I:%M %p ET")
ts = doc.add_paragraph(f"Report Run: {timestamp_str}")
ts.alignment = WD_ALIGN_PARAGRAPH.CENTER
ts.runs[0].font.size = Pt(12)
ts.runs[0].font.color.rgb = RGBColor(110, 110, 110)

doc.add_paragraph()

prepared = doc.add_paragraph("Prepared for: Tom Short")
prepared.alignment = WD_ALIGN_PARAGRAPH.CENTER
prepared.runs[0].font.size = Pt(13)

doc.add_page_break()

# Executive Summary
doc.add_heading("Executive Summary", level=1)

# Short intro paragraph
doc.add_paragraph(
    "This executive summary provides a high-level snapshot of your portfolio using live market data, "
    "including valuation, recent performance, diversification, rebalancing needs, and positioning versus key benchmarks."
)

# -------------------- Portfolio Snapshot (Table) --------------------
h_ps = doc.add_heading("Portfolio Snapshot", level=2)
h_ps.paragraph_format.space_before = Pt(6)

snapshot_rows = []
snapshot_rows.append(["Total Value", fmt_dollar(summary_total_value)])


if summary_target_value is not None:
    snapshot_rows.append(["Target Portfolio Value", fmt_dollar(summary_target_value)])

snapshot_rows.append(["MTD Return", fmt_pct(port_mtd) if not np.isnan(port_mtd) else "-"])
snapshot_rows.append(["YTD Return", fmt_pct(port_ytd) if not np.isnan(port_ytd) else "-"])
snapshot_rows.append(["Total 1M P/L", fmt_dollar(total_1m_pl)])
snapshot_rows.append(["Number of Holdings", str(summary_num_holdings)])

add_table(
    ["Metric", "Value"],
    snapshot_rows,
    right_align_cols=[1],
)

# ----------------- Performance Highlights (Table) -------------------
h_ph = doc.add_heading("Performance Highlights", level=2)
h_ph.paragraph_format.space_before = Pt(10)

perf_rows = [

    ["Top 1M Performer", top_1m_line],
    ["Bottom 1M Performer", bottom_1m_line],
    ["Best 1D Performer", best_1d_line if best_1d_line else "N/A"],
]

add_table(
    ["Metric", "Value"],
    perf_rows,
    right_align_cols=[1],
)

# --------------- Risk & Diversification (Table) ---------------------
h_rd = doc.add_heading("Risk & Diversification", level=2)
h_rd.paragraph_format.space_before = Pt(10)

top3_str = fmt_pct(top3_pct) if not np.isnan(top3_pct) else "-"
largest_ac_str = (
    f"{largest_ac_name} ({largest_ac_pct:.2f}%)" if not np.isnan(largest_ac_pct) else "N/A"
)

risk_rows = [
    ["Top 3 holdings % of portfolio", top3_str],
    ["Largest asset class", largest_ac_str],
    ["Largest overweight", largest_overweight_str],
    ["Largest underweight", largest_underweight_str],
]

add_table(
    ["Metric", "Value"],
    risk_rows,
    right_align_cols=[1],
)

# --------- Rebalancing & Benchmark Positioning (Table) --------------
h_rb = doc.add_heading("Rebalancing & Benchmark Positioning", level=2)
h_rb.paragraph_format.space_before = Pt(10)

pct_target_str = f"{pct_tickers_on_target:.1f}%" if pct_tickers_on_target is not None else "N/A"

reb_bench_rows = [
    ["Rebalancing need", fmt_dollar(rebalance_need)],
    ["% of tickers on target (±5% band)", pct_target_str],
    ["Vs S&P 500 MTD", vs_sp500_mtd_str],
    ["Vs S&P 500 YTD", vs_sp500_ytd_str],
]

add_table(
    ["Metric", "Value"],
    reb_bench_rows,
    right_align_cols=[1],
)

doc.add_page_break()


# Holdings by Ticker
doc.add_heading("Portfolio Composition & Strategy", level=1)
doc.add_heading("Holdings by Ticker", level=2)
ticker_rows = []
for _, row in df.iterrows():
    ticker_rows.append(
        [
            row["ticker"],
            row["asset_class"],
            f"{row['shares']:.4f}",
            f"{row['price']:.2f}",
            f"{row['value']:.2f}",
            f"{row['core_allocation_pct']:.2f}%",
            f"{(row.get('target_pct') or 0.0):.2f}%",
            "-"
            if pd.isna(row.get("contribute_to_target"))
            else fmt_dollar(row["contribute_to_target"]),
        ]
    )

total_value_sum = df["value"].sum()
total_contrib_sum = df["contribute_to_target"].fillna(0).sum()

ticker_rows.append(
    [
        "TOTAL",
        "",
        "",
        "",
        f"{total_value_sum:,.2f}",
        "",
        "",
        fmt_dollar(total_contrib_sum),
    ]
)

add_table(
    [
        "Ticker",
        "Asset Class",
        "Shares",
        "Price ($)",
        "Value ($)",
        "Allocation %",
        "Target %",
        "Contribute to Target ($)",
    ],
    ticker_rows,
    right_align_cols=[2, 3, 4, 5, 6, 7],
)

# --------- Illustrative Monthly Contribution Schedule (Table) ---------
if monthly_contrib > 0 and not df.empty:
    # Positive gaps only (how much each ticker needs to reach its target)
    need_series = df["contribute_to_target"].clip(lower=0).fillna(0)
    total_need_schedule = float(need_series.sum())

    if total_need_schedule > 0:
        sched_rows = []
        for _, row in df.iterrows():
            gap = float(row.get("contribute_to_target") or 0.0)
            if gap <= 0:
                continue

            share = gap / total_need_schedule        # fraction of total shortfall
            monthly_dollars = monthly_contrib * share

            sched_rows.append(
                [
                    row["ticker"],
                    row["asset_class"],
                    fmt_dollar(gap),
                    fmt_dollar(monthly_dollars),
                    f"{share * 100:.1f}%",
                ]
            )

        if sched_rows:
            # Approximate months of contributions to close current gaps,
            # assuming flat markets and this allocation split
            portfolio_months_to_target = total_need_schedule / monthly_contrib

            h_sched = doc.add_heading(
                "Illustrative Monthly Contribution Schedule", level=2
            )
            h_sched.paragraph_format.space_before = Pt(10)

            add_table(
                ["Ticker", "Asset Class", "Gap to Target ($)", "Suggested Monthly ($)", "Share of Monthly %"],
                sched_rows,
                right_align_cols=[2, 3, 4],
            )

            p_sched_note = doc.add_paragraph(
                f"At approximately ${monthly_contrib:,.0f}/month, this schedule would allocate "
                f"contributions proportionally to each holding's shortfall. In total, it would "
                f"take about {portfolio_months_to_target:.1f} months of contributions to fully "
                "close the current gaps, assuming markets are flat and this allocation is followed."
            )
            p_sched_note.paragraph_format.space_before = Pt(4)


# Asset Class Allocation Overview
h_ac = doc.add_heading("Asset Class Allocation Overview", level=2)
h_ac.paragraph_format.space_before = Pt(12)

if asset_targets_df is not None and not asset_targets_df.empty:
    ac_compare = asset_df.merge(
        asset_targets_df, on="asset_class", how="left", suffixes=("", "_target")
    )
    ac_compare["target_pct"] = ac_compare["target_pct"].fillna(0.0)
    ac_compare["delta_pct"] = ac_compare["allocation_pct"] - ac_compare["target_pct"]

    ac_compare = ac_compare.reindex(
        ac_compare["delta_pct"].abs().sort_values(ascending=False).index
    )

    ac_rows = []
    for _, row in ac_compare.iterrows():
        delta_str = fmt_pct(row["delta_pct"])
        ac_rows.append(
            [
                row["asset_class"],
                f"{row['value']:.2f}",
                f"{row['allocation_pct']:.2f}%",
                f"{row['target_pct']:.2f}%",
                delta_str,
            ]
        )

    total_ac_value = asset_df["value"].sum()

    ac_rows.append(
        [
            "TOTAL",
            f"{total_ac_value:,.2f}",
            "100.00%",
            f"{asset_targets_df['target_pct'].sum():.2f}%"
            if (asset_targets_df is not None and not asset_targets_df.empty)
            else "",
            "",
        ]
    )

    add_table(
        ["Asset Class", "Value ($)", "Actual %", "Target %", "Delta %"],
        ac_rows,
        right_align_cols=[1, 2, 3, 4],
    )

else:
    ac_rows = []
    for _, row in asset_df.iterrows():
        ac_rows.append(
            [
                row["asset_class"],
                f"{row['value']:.2f}",
                f"{row['allocation_pct']:.2f}%",
            ]
        )
    add_table(
        ["Asset Class", "Value ($)", "Allocation %"],
        ac_rows,
        right_align_cols=[1, 2],
    )


# Next Steps & Ongoing Strategy
h_ns = doc.add_heading("Next Steps & Ongoing Strategy", level=2)
h_ns.paragraph_format.space_before = Pt(12)

add_table(
    ["Focus Area", "Guidance"],
    [
        ["Rebalancing", "Review annually; adjust if allocations drift ±5%."],
        [
            "Contributions",
            f"Automate around ${monthly_contrib:,.0f}/mo into your target allocation.",
        ],
        [
            "Risk Management",
            "Stay within your risk tolerance; increase bonds as you approach major goals.",
        ],
        [
            "Monitoring",
            "Avoid overreacting to short-term volatility; review quarterly or annually.",
        ],
    ],
)


# Visuals – Allocation & Diversification
doc.add_page_break()
doc.add_heading("Visual Report – Allocation & Diversification", level=1)

doc.add_heading("Ticker-Level Allocation Breakdown", level=2)
doc.add_picture(ticker_pie_stream, width=Inches(5.5))
p = doc.add_paragraph("Figure 1: Allocation by ticker.")
p.alignment = WD_ALIGN_PARAGRAPH.CENTER

doc.add_paragraph()
doc.add_heading("Asset Class Allocation Breakdown", level=2)
doc.add_picture(asset_pie_stream, width=Inches(5.5))
p = doc.add_paragraph("Figure 2: Allocation by asset class.")
p.alignment = WD_ALIGN_PARAGRAPH.CENTER

doc.add_paragraph()
doc.add_heading("Sector Allocation Heatmap", level=2)
doc.add_picture(sector_stream, width=Inches(5.5))
p = doc.add_paragraph(
    "Figure 3: Approximate sector exposure based on underlying holdings."
)
p.alignment = WD_ALIGN_PARAGRAPH.CENTER

# Performance & Growth
doc.add_page_break()
doc.add_heading("Performance & Growth", level=1)
doc.add_heading("Performance vs Benchmarks (MTD & YTD)", level=2)

bench_df["MTD %"] = pd.to_numeric(bench_df["MTD %"], errors="coerce")
bench_df["YTD %"] = pd.to_numeric(bench_df["YTD %"], errors="coerce")
if "Portfolio (Live)" in bench_df["Benchmark"].values:
    port_row = bench_df.loc[bench_df["Benchmark"] == "Portfolio (Live)"].iloc[0]
else:
    port_row = pd.Series({"MTD %": np.nan, "YTD %": np.nan})

mtd_rows, ytd_rows = [], []
for _, row in bench_df.iterrows():
    if row["Benchmark"] == "Portfolio (Live)":
        continue

    port_mtd_val = port_row.get("MTD %")
    bench_mtd_val = row["MTD %"]
    if pd.isna(port_mtd_val) or pd.isna(bench_mtd_val):
        excess_mtd = None
    else:
        excess_mtd = float(port_mtd_val) - float(bench_mtd_val)

    port_ytd_val = port_row.get("YTD %")
    bench_ytd_val = row["YTD %"]
    if pd.isna(port_ytd_val) or pd.isna(bench_ytd_val):
        excess_ytd = None
    else:
        excess_ytd = float(port_ytd_val) - float(bench_ytd_val)

    mtd_rows.append(
        [
            row["Benchmark"],
            fmt_pct(port_mtd_val),
            fmt_pct(bench_mtd_val),
            fmt_pct(excess_mtd),
        ]
    )
    ytd_rows.append(
        [
            row["Benchmark"],
            fmt_pct(port_ytd_val),
            fmt_pct(bench_ytd_val),
            fmt_pct(excess_ytd),
        ]
    )

# --- MTD Table ---
p_mtd = doc.add_paragraph("Month-to-date (MTD) comparison:")
p_mtd.paragraph_format.space_before = Pt(6)

add_table(
    ["Benchmark", "Portfolio MTD %", "Benchmark MTD %", "Excess MTD %"],
    mtd_rows,
    right_align_cols=[1, 2, 3],
)

# --- YTD Table ---
p_ytd = doc.add_paragraph("Year-to-date (YTD) comparison:")
p_ytd.paragraph_format.space_before = Pt(12)  # <-- give breathing room

add_table(
    ["Benchmark", "Portfolio YTD %", "Benchmark YTD %", "Excess YTD %"],
    ytd_rows,
    right_align_cols=[1, 2, 3],
)


doc.add_paragraph()
doc.add_heading("Holdings Multi-Horizon Returns", level=2)
doc.add_paragraph(
    "Performance of each holding over multiple lookback periods, ranked by 6-month return:"
)

returns_sorted = returns_df.sort_values("6M %", ascending=False, na_position="last")
rows = []
for _, r in returns_sorted.iterrows():
    rows.append(
        [
            r["Ticker"],
            fmt_pct(r["1D %"]),
            fmt_pct(r["1W %"]),
            fmt_pct(r["1M %"]),
            fmt_pct(r["3M %"]),
            fmt_pct(r["6M %"]),
        ]
    )

add_table(
    ["Ticker", "1D %", "1W %", "1M %", "3M %", "6M %"],
    rows,
    right_align_cols=[1, 2, 3, 4, 5],
)

# ultra-tight spacer below the returns table (no extra vertical gap)
sp = doc.add_paragraph()
sp.paragraph_format.space_before = Pt(0)
sp.paragraph_format.space_after = Pt(0)

# Dollar Profit/Loss Table
dollar_pl_sorted = (
    dollar_pl_df.set_index("Ticker").reindex(returns_sorted["Ticker"]).reset_index()
)

# Heading for P/L section – minimal spacing so it sits close to returns table
p_title = doc.add_heading("Holdings Multi-Horizon Profit/Loss ($)", level=2)
p_title.paragraph_format.space_before = Pt(2)
p_title.paragraph_format.keep_with_next = True

p_desc = doc.add_paragraph(
    "Approximate dollar profit/loss for each holding over the same lookback windows, "
    "using the same return calculations as the percentage table above."
)
p_desc.paragraph_format.space_before = Pt(0)
p_desc.paragraph_format.keep_with_next = True


pl_rows = []
for _, r in dollar_pl_sorted.iterrows():
    pl_rows.append(
        [
            r["Ticker"],
            fmt_dollar(r["1D $"]),
            fmt_dollar(r["1W $"]),
            fmt_dollar(r["1M $"]),
            fmt_dollar(r["3M $"]),
            fmt_dollar(r["6M $"]),
        ]
    )

add_table(
    ["Ticker", "1D $", "1W $", "1M $", "3M $", "6M $"],
    pl_rows,
    right_align_cols=[1, 2, 3, 4, 5],
)

doc.add_paragraph()
doc.add_heading("Compound Value Breakdown", level=2)
doc.add_picture(compound_stream, width=Inches(6))

# Shorter, tighter caption so the chart can be taller on the page
p = doc.add_paragraph("Figure 4: Illustrates how monthly contributions and compound market growth each drive total portfolio value over 20 years at a 7% annual return.")
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_before = Pt(2)
p.paragraph_format.space_after = Pt(4)

# 👉 New: start a fresh page for all remaining visuals together
doc.add_page_break()


doc.add_heading("20-Year Projection Scenarios", level=2)
contrib_label = f"+${int(monthly_contrib):,}/mo"
proj_headers = [
    "Year",
    "5% (Lump)",
    "7% (Lump)",
    "9% (Lump)",
    f"5% ({contrib_label})",
    f"7% ({contrib_label})",
    f"9% ({contrib_label})",
]

proj_rows_for_table = []
for row in proj_rows:
    formatted_row = [str(row[0])]
    for i in range(1, len(proj_headers)):
        formatted_row.append(f"{row[i]:,.0f}")
    proj_rows_for_table.append(formatted_row)

add_table(
    proj_headers,
    proj_rows_for_table,
    right_align_cols=list(range(1, len(proj_headers))),
)

doc.add_paragraph()
doc.add_picture(growth_stream, width=Inches(6))
p = doc.add_paragraph(
    "Figure 5: Variable long-term projections."
)
p.alignment = WD_ALIGN_PARAGRAPH.CENTER

# Risk & Volatility  (no page break here anymore)
doc.add_heading("Risk & Volatility Analysis", level=1)

# --- Risk Charts (Wider + Shorter + Tight Spacing) ---

h_vol = doc.add_heading("Expected Volatility by Asset Class", level=2)
h_vol.paragraph_format.space_before = Pt(4)
h_vol.paragraph_format.space_after = Pt(2)

doc.add_picture(vol_stream, width=Inches(5.25), height=Inches(3.4))
p = doc.add_paragraph(
    "Figure 6: Approximate volatility (standard deviation) by asset class."
)
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_before = Pt(1)
p.paragraph_format.space_after = Pt(12)

h_risk = doc.add_heading("Risk vs Expected Return", level=2)
h_risk.paragraph_format.space_before = Pt(4)
h_risk.paragraph_format.space_after = Pt(2)

doc.add_picture(risk_stream, width=Inches(5.25), height=Inches(3.4))
p = doc.add_paragraph(
    "Figure 7: Trade-off between expected return and volatility by asset class."
)
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_before = Pt(1)
p.paragraph_format.space_after = Pt(0)


# ------------------------- 13) SAVE DOCX + PDF ------------------------

today_str = get_eastern_now().strftime("%Y-%m-%d")
docx_name = f"Investment_Report_{today_str}.docx"
pdf_name = f"Investment_Report_{today_str}.pdf"

doc.save(docx_name)
print(f"Report generated: {docx_name}")

if os.path.exists(pdf_name):
    os.remove(pdf_name)

if os.path.exists(docx_name):
    os.utime(docx_name, None)

pdf_created = False

# Only generate PDF on Windows / macOS; skip on Linux (Colab)
if sys.platform.startswith("win") or sys.platform == "darwin":
    try:
        from docx2pdf import convert  # lazy import so Colab / Linux don't care
        print("Attempting PDF export via docx2pdf (Word)...")
        convert(docx_name, pdf_name)
        if os.path.exists(pdf_name):
            pdf_created = True
            print(f"✔ PDF created via docx2pdf: {pdf_name}")
        else:
            print("⚠ docx2pdf ran, but PDF file not found.")
    except Exception as e:
        print(f"❌ PDF export failed via docx2pdf: {e}")
        print("Make sure Microsoft Word is installed and docx2pdf is configured.")
else:
    print("Skipping PDF generation on this platform (DOCX only).")

if not pdf_created:
    print("⚠ No PDF generated. You still have the DOCX file.")

# Copy outputs into Drive / other folders
try:
    dest_dirs = list(EXTRA_OUTPUT_DIRS) if "EXTRA_OUTPUT_DIRS" in globals() else []

    colab_drive_outputs = "/content/drive/MyDrive/Investment Report Outputs"
    if os.path.isdir(colab_drive_outputs) and colab_drive_outputs not in dest_dirs:
        dest_dirs.append(colab_drive_outputs)

    files_to_copy = [docx_name]
    if pdf_created and os.path.exists(pdf_name):
        files_to_copy.append(pdf_name)

    for out_dir in dest_dirs:
        if not os.path.isdir(out_dir):
            continue
        for fname in files_to_copy:
            src = os.path.abspath(fname)
            dst = os.path.join(out_dir, os.path.basename(fname))
            try:
                shutil.copy2(src, dst)
                print(f"Copied {fname} -> {dst}")
            except Exception as copy_err:
                print(f"Could not copy {fname} to {out_dir}: {copy_err}")

    if dest_dirs:
        if pdf_created:
            print("✔ Copied DOCX and PDF to output folders.")
        else:
            print("✔ Copied DOCX to output folders (no PDF).")
except Exception as outer_err:
    print(f"Post-processing copy step failed: {outer_err}")
