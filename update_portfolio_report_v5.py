# ==============================================================
# update_portfolio_report_v3.py  (Part 1 of 2)
# Live portfolio report with allocation, diversification,
# and real-time MTD / YTD benchmark comparisons
# ==============================================================

import yfinance as yf
import pandas as pd
import numpy as np
from io import BytesIO
import matplotlib.pyplot as plt
from datetime import datetime
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx2pdf import convert

def fmt_pct(x):
    import math
    if x is None or (isinstance(x, float) and math.isnan(x)):
        return "-"
    return f"{x:+.2f}%"

def fmt_dollar(x):
    import math
    if x is None or (isinstance(x, float) and math.isnan(x)):
        return "-"
    return f"${x:,.2f}"

# --------------------------------------------------------------
# CONFIG
# --------------------------------------------------------------
HOLDINGS_CSV = "holdings.csv"
RISK_FREE_RATE = 0.04     # for Sharpe ratio later
monthly_contrib = 500.0
COLOR_MAIN = ["#2563EB", "#10B981", "#F59E0B", "#6366F1", "#14B8A6"]
# --------------------------------------------------------------
# TARGET ALLOCATIONS — Core Portfolio Only
# --------------------------------------------------------------
TICKER_TARGETS_PCT = {
    "VOO": 44.0,
    "QQQ": 11.0,
    "VXUS": 15.0,
    "BND": 15.0,
    "GLD": 10.0,
    "FBTC": 5.0,
}

# Fixed holdings (excluded from target allocation math)
FIXED_TICKERS = ["AMZN", "NDAQ"]


# --------------------------------------------------------------
# 1️⃣  LOAD HOLDINGS  +  LIVE PRICES
# --------------------------------------------------------------
df = pd.read_csv(HOLDINGS_CSV)

def get_live_price(ticker):
    data = yf.Ticker(ticker).history(period="1d")
    if data.empty:
        raise ValueError(f"No price data for {ticker}")
    return float(data["Close"].iloc[-1])

prices = [get_live_price(t) for t in df["ticker"]]
df["price"] = prices
df["value"] = df["shares"] * df["price"]
total_value = df["value"].sum()
df["allocation_pct"] = (df["value"] / total_value * 100).round(2)

# --------------------------------------------------------------
# 2️⃣   ASSET-CLASS  &  NORMALIZATION
# --------------------------------------------------------------
asset_df = df.groupby("asset_class")["value"].sum().reset_index()
asset_df["allocation_pct"] = (asset_df["value"] / total_value * 100).round(2)

def normalize_allocations(series):
    rounded = series.round(2)
    diff = 100.00 - rounded.sum()
    rounded.iloc[-1] += diff
    return rounded

df["allocation_pct"] = normalize_allocations(df["allocation_pct"])
asset_df["allocation_pct"] = normalize_allocations(asset_df["allocation_pct"])


# --------------------------------------------------------------
# PER-TICKER TARGETS (core portfolio only, fixed tickers excluded)
# --------------------------------------------------------------
# Mark fixed tickers
df["is_fixed"] = df["ticker"].isin(FIXED_TICKERS)

# Totals
total_value = df["value"].sum()
fixed_value = df.loc[df["is_fixed"], "value"].sum()
allocatable_total = total_value - fixed_value

# Apply targets to core tickers only
core_df = df[~df["is_fixed"]].copy()

# ----- Core-only snapshots (exclude AMZN/NDAQ) -----
core_total = allocatable_total
core_df["core_allocation_pct"] = (core_df["value"] / core_total * 100).round(2)

# Core-only asset-class summary
core_asset_df = core_df.groupby("asset_class", as_index=False)["value"].sum()
core_asset_df["allocation_pct"] = (core_asset_df["value"] / core_total * 100).round(2)

# Map % targets to core tickers
core_df["target_pct_raw"] = core_df["ticker"].map(TICKER_TARGETS_PCT).astype(float)

# Dollar targets for core = % of allocatable_total
core_df["target_value"] = allocatable_total * (core_df["target_pct_raw"] / 100.0)

# Display Target % as share of the core portfolio only
core_df["target_pct"] = core_df["target_value"] / allocatable_total * 100.0

# Merge back into main df
df = df.merge(core_df[["ticker", "target_value", "target_pct", "core_allocation_pct"]], on="ticker", how="left")

# Contribute-only dollars (no selling)
df["delta_to_target_raw"] = df["target_value"] - df["value"]
df["contribute_to_target"] = np.where(
    df["delta_to_target_raw"].isna(), np.nan,
    np.where(df["delta_to_target_raw"] > 0, df["delta_to_target_raw"], 0.0)
)


# --------------------------------------------------------------
# 3️⃣  SECTOR  &  GEOGRAPHIC  APPROXIMATIONS
# --------------------------------------------------------------
# Hard-coded estimated weights for simplicity
sector_weights = {
    "Information Technology": 32,
    "Financials": 12,
    "Healthcare": 11,
    "Industrials": 9,
    "Consumer Discretionary": 10,
    "Communication Services": 8,
    "Energy": 4,
    "Materials": 3,
    "Real Estate": 3,
    "Utilities": 3,
    "Bitcoin / Digital Assets": 3,
    "Gold / Precious Metals": 2,
}
sector_df = pd.DataFrame(
    {"Sector": list(sector_weights.keys()), "Weight": list(sector_weights.values())}
)

geo_alloc = {
    "United States": 70,
    "International Developed": 20,
    "Emerging Markets": 5,
    "Global Bonds": 3,
    "Precious Metals": 2,
}
geo_df = pd.DataFrame(
    {"Region": list(geo_alloc.keys()), "Weight": list(geo_alloc.values())}
)

# --------------------------------------------------------------
# 4️⃣  SECTOR HEATMAP (Bar Chart)
# --------------------------------------------------------------
plt.figure(figsize=(6, 5))
bars = plt.barh(
    sector_df["Sector"], sector_df["Weight"],
    color=plt.cm.Blues(np.linspace(0.4, 0.9, len(sector_df)))
)
plt.xlabel("Portfolio Exposure (%)")
plt.title("Sector Allocation Heatmap", fontsize=12, weight="bold")
plt.gca().invert_yaxis()
plt.tight_layout()
sector_stream = BytesIO()
plt.savefig(sector_stream, format="png", bbox_inches="tight", facecolor="white")
sector_stream.seek(0)
plt.close()


# --------------------------------------------------------------
# 6️⃣  BENCHMARK DATA (MTD & YTD)
# --------------------------------------------------------------

def get_return_pct(ticker, start_date, end_date):
    """Fetch percentage return (in %) for a ticker between start_date and end_date."""
    try:
        data = yf.download(ticker, start=start_date, end=end_date, progress=False)
        if data.empty:
            # fallback: look back a bit more if needed
            data = yf.download(ticker, period="60d", progress=False)
        if data.empty:
            return np.nan

        close = data["Close"]
        end = float(close.iloc[-1])
        start = float(close.iloc[0])
        return (end / start - 1.0) * 100.0

    except Exception:
        return np.nan



# Real benchmark tickers
benchmarks = {
    "S&P 500": "^GSPC",
    "Global 60/40": "AOR",
    "Conservative 40/60": "AOK",
}

today = datetime.now()
# Use the last trading day of the previous month for MTD start
start_month = today.replace(day=1) - pd.tseries.offsets.BDay(1)
start_year = datetime(today.year, 1, 1) - pd.tseries.offsets.BDay(1)

bench_rows = []

# Benchmark MTD/YTD
for name, ticker in benchmarks.items():
    mtd = get_return_pct(ticker, start_month, today)
    ytd = get_return_pct(ticker, start_year, today)
    bench_rows.append(
        {
            "Benchmark": name,
            "MTD %": round(mtd, 2) if not np.isnan(mtd) else np.nan,
            "YTD %": round(ytd, 2) if not np.isnan(ytd) else np.nan,
        }
    )

# Portfolio MTD/YTD as allocation-weighted average of its holdings
port_changes_mtd = []
port_changes_ytd = []
weights = []

for _, row in df.iterrows():
    t = row["ticker"]
    w = row["allocation_pct"]

    mtd = get_return_pct(t, start_month, today)
    ytd = get_return_pct(t, start_year, today)

    if np.isnan(mtd) or np.isnan(ytd):
        continue

    port_changes_mtd.append(mtd)
    port_changes_ytd.append(ytd)
    weights.append(w)

def weighted_avg(values, weights):
    if not values:
        return np.nan
    v = np.array(values, dtype=float)
    w = np.array(weights, dtype=float)
    return float((v * w).sum() / w.sum())

port_mtd = weighted_avg(port_changes_mtd, weights)
port_ytd = weighted_avg(port_changes_ytd, weights)

bench_rows.insert(
    0,
    {
        "Benchmark": "Portfolio (Live)",
        "MTD %": round(port_mtd, 2) if not np.isnan(port_mtd) else np.nan,
        "YTD %": round(port_ytd, 2) if not np.isnan(port_ytd) else np.nan,
    },
)

bench_df = pd.DataFrame(bench_rows)

# --------------------------------------------------------------
# 6b️⃣  HOLDINGS MULTI-HORIZON RETURNS (1W, 1M, 3M, 6M)
# --------------------------------------------------------------

# Define horizons in days
horizons = {
    "1W %": 7,
    "1M %": 30,
    "3M %": 90,
    "6M %": 180,
}

returns_rows = []
for _, row in df.iterrows():
    t = row["ticker"]
    r = {"Ticker": t}
    for label, days in horizons.items():
        start = today - pd.Timedelta(days=days)
        val = get_return_pct(t, start, today)
        r[label] = round(val, 2) if not np.isnan(val) else np.nan
    returns_rows.append(r)

returns_df = pd.DataFrame(returns_rows)

# --------------------------------------------------------------
# 7️⃣  LONG-TERM PROJECTIONS (20 YEARS)
# --------------------------------------------------------------
rates = [0.05, 0.07, 0.09]
years = [1, 5, 10, 15, 20]

def future_value(principal, rate, years):
    return principal * ((1 + rate/12) ** (years * 12))

def future_value_with_contrib(principal, rate, years, monthly):
    months = years * 12
    monthly_rate = rate / 12
    fv_principal = principal * ((1 + monthly_rate) ** months)
    fv_contrib = monthly * (((1 + monthly_rate) ** months - 1) / monthly_rate)
    return fv_principal + fv_contrib

proj_rows = []
for y in years:
    row = [y]
    for r in rates:
        row.append(round(future_value(total_value, r, y)))
    for r in rates:
        row.append(round(future_value_with_contrib(total_value, r, y, monthly_contrib)))
    proj_rows.append(row)

# 20-year growth scenarios line chart (same as before)
growth_lump = [[future_value(total_value, r, y) for y in years] for r in rates]
growth_plus = [[future_value_with_contrib(total_value, r, y, monthly_contrib) for y in years] for r in rates]

plt.figure(figsize=(6, 4))
for i, r in enumerate(rates):
    plt.plot(years, growth_lump[i], label=f"{int(r*100)}% Lump Sum", linestyle='-')
    plt.plot(years, growth_plus[i], label=f"{int(r*100)}% + $500/mo", linestyle='--')
plt.title("Portfolio Growth Projections (20-Year Scenarios)", fontsize=12, weight="bold")
plt.xlabel("Years")
plt.ylabel("Portfolio Value ($)")
plt.grid(alpha=0.3)
plt.legend(fontsize=8)
plt.tight_layout()
growth_stream = BytesIO()
plt.savefig(growth_stream, format="png", bbox_inches="tight", facecolor="white")
growth_stream.seek(0)
plt.close()

# --------------------------------------------------------------
# 8️⃣  COMPOUND VALUE BREAKDOWN (CONTRIBUTIONS VS GROWTH)
# --------------------------------------------------------------
years_compound = list(range(0, 21))  # 0–20 years
rate_for_compound = 0.07

contrib_values = []
growth_values = []

for y in years_compound:
    total_with_contrib = future_value_with_contrib(total_value, rate_for_compound, y, monthly_contrib)
    total_contrib = monthly_contrib * 12 * y
    contrib_values.append(total_contrib)
    growth_values.append(max(total_with_contrib - total_contrib, 0))

plt.figure(figsize=(6, 4))
plt.stackplot(
    years_compound,
    contrib_values,
    growth_values,
    labels=["Contributions", "Growth"],
    colors=[COLOR_MAIN[0], COLOR_MAIN[1]],
    alpha=0.85
)
plt.title("Compound Value Breakdown (Contributions vs Growth)", fontsize=12, weight="bold")
plt.xlabel("Years")
plt.ylabel("Portfolio Value ($)")
plt.legend(loc="upper left", fontsize=8)
plt.tight_layout()
compound_stream = BytesIO()
plt.savefig(compound_stream, format="png", bbox_inches="tight", facecolor="white")
compound_stream.seek(0)
plt.close()

# 9️⃣  PERFORMANCE VS BENCHMARKS (MTD & YTD)
# --------------------------------------------------------------
# Clean numeric data for plotting
bench_df["MTD %"] = pd.to_numeric(bench_df["MTD %"], errors="coerce")
bench_df["YTD %"] = pd.to_numeric(bench_df["YTD %"], errors="coerce")

# Drop any rows where both MTD and YTD are NaN
bench_plot = bench_df.dropna(subset=["MTD %", "YTD %"], how="all").reset_index(drop=True)

# If everything is missing, just make a blank info chart
if bench_plot.empty:
    plt.figure(figsize=(6, 4))
    plt.text(
        0.5, 0.5,
        "Benchmark return data unavailable.\n(Check internet connection or market trading days.)",
        ha="center", va="center", fontsize=9,
    )
    plt.axis("off")
    plt.tight_layout()
else:
    x = np.arange(len(bench_plot))  # Portfolio + benchmarks
    width = 0.35

    mtd_vals = bench_plot["MTD %"].to_numpy(dtype=float)
    ytd_vals = bench_plot["YTD %"].to_numpy(dtype=float)

    plt.figure(figsize=(6, 4))
    mtd_bars = plt.bar(x - width/2, mtd_vals, width, label="MTD", color=COLOR_MAIN[0])
    ytd_bars = plt.bar(x + width/2, ytd_vals, width, label="YTD", color=COLOR_MAIN[1])

    plt.xticks(x, bench_plot["Benchmark"], rotation=20, ha="right")
    plt.ylabel("Return (%)")
    plt.title("Portfolio vs Benchmarks (MTD & YTD)", fontsize=12, weight="bold")
    plt.legend(fontsize=8)
    plt.grid(axis="y", alpha=0.3)

    # Dynamic y-limits so small returns are visible
    all_vals = np.concatenate([mtd_vals[~np.isnan(mtd_vals)], ytd_vals[~np.isnan(ytd_vals)]])
    if len(all_vals) > 0:
        vmin, vmax = float(all_vals.min()), float(all_vals.max())
        if vmin == vmax:  # all same value, give some padding
            vmin -= 1.0
            vmax += 1.0
        else:
            pad = (vmax - vmin) * 0.2
            vmin -= pad
            vmax += pad
        plt.ylim(vmin, vmax)

    # Label each bar with its value
    def label_bars(bars, vals):
        for bar, val in zip(bars, vals):
            if np.isnan(val):
                continue
            height = bar.get_height()
            plt.text(
                bar.get_x() + bar.get_width()/2,
                height,
                f"{val:+.1f}%",
                ha="center",
                va="bottom",
                fontsize=7,
            )

    label_bars(mtd_bars, mtd_vals)
    label_bars(ytd_bars, ytd_vals)

    # annotate excess YTD vs portfolio on YTD bars for benchmarks only
    if "Portfolio (Live)" in bench_plot["Benchmark"].values:
        port_ytd_val = bench_plot.loc[
            bench_plot["Benchmark"] == "Portfolio (Live)", "YTD %"
        ].iloc[0]
        for i, row in bench_plot.iterrows():
            if row["Benchmark"] == "Portfolio (Live)":
                continue
            if pd.isna(row["YTD %"]) or pd.isna(port_ytd_val):
                continue
            excess = port_ytd_val - row["YTD %"]
            plt.text(
                x[i] + width / 2,
                ytd_vals[i] + (vmax - vmin) * 0.05,
                f"Excess {excess:+.1f}%",
                ha="center",
                va="bottom",
                fontsize=7,
            )

    plt.tight_layout()


# --------------------------------------------------------------
# 🔟  RISK & VOLATILITY (SAME IDEA AS BEFORE)
# --------------------------------------------------------------
RISK_RETURN = {
    "US Equities":            {"return": 8.0,  "vol": 15.0},
    "International Equities": {"return": 8.5,  "vol": 17.0},
    "Fixed Income":           {"return": 4.0,  "vol": 5.0},
    "Precious Metals":        {"return": 6.0,  "vol": 10.0},
    "Digital Assets":         {"return": 15.0, "vol": 40.0},
}

# Use core-only asset weights for risk visuals
risk_rows = []
for _, row in core_asset_df.iterrows():
    ac = row["asset_class"]
    info = RISK_RETURN.get(ac, {"return": 0.0, "vol": 0.0})
    risk_rows.append({
        "asset_class": ac,
        "vol": info["vol"],
        "ret": info["return"]
    })


risk_df = pd.DataFrame(risk_rows)

# Volatility bar chart
plt.figure(figsize=(6, 4))
plt.bar(risk_df["asset_class"], risk_df["vol"], color=COLOR_MAIN[0])
plt.title("Expected Volatility by Asset Class", fontsize=12, weight="bold")
plt.ylabel("Standard Deviation (%)")
plt.xticks(rotation=25, ha="right")
plt.tight_layout()
vol_stream = BytesIO()
plt.savefig(vol_stream, format="png", bbox_inches="tight", facecolor="white")
vol_stream.seek(0)
plt.close()

# Risk vs Return scatter
plt.figure(figsize=(6, 4))
plt.scatter(risk_df["vol"], risk_df["ret"], s=80, color=COLOR_MAIN[1])
for i, row in risk_df.iterrows():
    plt.annotate(row["asset_class"], (row["vol"] + 0.4, row["ret"]), fontsize=8, weight="bold")
plt.title("Risk vs Expected Return by Asset Class", fontsize=12, weight="bold")
plt.xlabel("Volatility (Std Dev %)")
plt.ylabel("Expected Annual Return (%)")
plt.grid(alpha=0.3)
plt.tight_layout()
risk_stream = BytesIO()
plt.savefig(risk_stream, format="png", bbox_inches="tight", facecolor="white")
risk_stream.seek(0)
plt.close()

# 1️⃣1️⃣  TICKER & ASSET ALLOCATION PIES (CORE ONLY)
# Ticker pie (exclude AMZN/NDAQ)
plt.figure(figsize=(6, 6))
plt.pie(
    core_df["core_allocation_pct"],
    labels=core_df["ticker"],
    autopct="%1.2f%%",
    startangle=90,
    pctdistance=0.85,
    labeldistance=1.05,
)
plt.title("Portfolio Allocation by Ticker — Core Only (excludes AMZN & NDAQ)", fontsize=12, weight="bold")
plt.tight_layout()
ticker_pie_stream = BytesIO()
plt.savefig(ticker_pie_stream, format="png", bbox_inches="tight", facecolor="white")
ticker_pie_stream.seek(0)
plt.close()

# Asset class pie (exclude AMZN/NDAQ)
plt.figure(figsize=(6, 6))
plt.pie(
    core_asset_df["allocation_pct"],
    labels=core_asset_df["asset_class"],
    autopct="%1.2f%%",
    startangle=90,
    pctdistance=0.85,
    labeldistance=1.05,
)
plt.title("Asset Class Allocation — Core Only (excludes AMZN & NDAQ)", fontsize=12, weight="bold")
plt.tight_layout()
asset_pie_stream = BytesIO()
plt.savefig(asset_pie_stream, format="png", bbox_inches="tight", facecolor="white")
asset_pie_stream.seek(0)
plt.close()

# --------------------------------------------------------------
# 1️⃣2️⃣  BUILD WORD DOCUMENT
# --------------------------------------------------------------
doc = Document()
style = doc.styles["Normal"]
style.font.name = "Calibri"
style._element.rPr.rFonts.set(qn("w:eastAsia"), "Calibri")
style.font.size = Pt(11)

def add_table(headers, rows, right_align_cols=None):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Light Grid Accent 1"
    hdr_cells = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr_cells[i].text = h
        for p in hdr_cells[i].paragraphs:
            for r in p.runs:
                r.bold = True
    for row_data in rows:
        row_cells = table.add_row().cells
        for i, val in enumerate(row_data):
            row_cells[i].text = str(val)
    if right_align_cols is not None:
        for row in table.rows[1:]:
            for idx in right_align_cols:
                cell = row.cells[idx]
                for p in cell.paragraphs:
                    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    return table

# Cover
doc.add_paragraph("\n\n\n\n\n")
cover = doc.add_paragraph("Comprehensive Investment Report — Live Portfolio")
cover.runs[0].bold = True
cover.runs[0].font.size = Pt(26)
cover.alignment = WD_ALIGN_PARAGRAPH.CENTER

subtitle = doc.add_paragraph("Automated Summary, Diversification, and Performance vs Benchmarks")
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
subtitle.runs[0].font.size = Pt(14)
subtitle.runs[0].font.color.rgb = RGBColor(70, 130, 180)

doc.add_paragraph()
doc.add_paragraph("Prepared for: [Your Name Here]").alignment = WD_ALIGN_PARAGRAPH.CENTER
doc.add_page_break()

# Executive Summary
doc.add_heading("Executive Summary", level=1)
doc.add_paragraph(
    f"This report summarizes your current portfolio based on live market data. "
    f"Total portfolio value is approximately ${total_value:,.0f}. "
    f"Allocations, diversification, and performance vs benchmarks are calculated using your most recent holdings."
)

# Overall Summary
doc.add_heading("Overall Summary", level=1)
overall_rows = [
    ["Total Portfolio Value", f"{total_value:,.2f}", "Sum of all holdings at live prices"],
]
add_table(["Category", "Amount ($)", "Notes"], overall_rows, right_align_cols=[1])

# Holdings by Ticker (target % + contribute-only $)
# Holdings by Ticker (target % + contribute-only $)
doc.add_heading("Holdings by Ticker", level=1)

# Add note about fixed holdings
note = doc.add_paragraph(
    "Note: AMZN and NDAQ are excluded from target allocations and remain fixed holdings."
)
note.runs[0].italic = True
note.alignment = WD_ALIGN_PARAGRAPH.CENTER

ticker_rows = []

for _, row in df.iterrows():
    ticker_rows.append([
        row["ticker"],
        row["asset_class"],
        f"{row['shares']:.4f}",
        f"{row['price']:.2f}",
        f"{row['value']:.2f}",
        "-" if row.get("is_fixed") else f"{row['core_allocation_pct']:.2f}%",
        "-" if pd.isna(row.get("target_pct")) else f"{row['target_pct']:.2f}%",
        "-" if pd.isna(row.get("contribute_to_target")) else fmt_dollar(row["contribute_to_target"]),
    ])

add_table(
    ["Ticker", "Asset Class", "Shares", "Price ($)", "Value ($)", "Allocation %", "Target %", "Contribute to Target ($)"],
    ticker_rows,
    right_align_cols=[2, 3, 4, 5, 6, 7]
)


# Asset Class Allocation Overview (no targets)
doc.add_heading("Asset Class Allocation Overview", level=1)
ac_rows = []
for _, row in asset_df.iterrows():
    ac_rows.append([
        row["asset_class"],
        f"{row['value']:.2f}",
        f"{row['allocation_pct']:.2f}%",
    ])

add_table(
    ["Asset Class", "Value ($)", "Allocation %"],
    ac_rows,
    right_align_cols=[1, 2]
)


# Next Steps & Ongoing Strategy
doc.add_page_break()  # <<< add this line
doc.add_heading("Next Steps & Ongoing Strategy", level=1)
add_table(
    ["Focus Area", "Guidance"],
    [
        ["Rebalancing", "Review annually; adjust if allocations drift ±5%."],
        ["Contributions", f"Automate around ${monthly_contrib:,.0f}/mo into your target allocation."],
        ["Risk Management", "Stay within your risk tolerance; increase bonds as you approach major goals."],
        ["Monitoring", "Avoid overreacting to short-term volatility; review quarterly or annually."],
    ],
)


# Totals Check
doc.add_heading("Totals Check", level=1)
add_table(
    ["Metric", "Value"],
    [
        ["Total Portfolio Value", f"{total_value:,.2f}"],
    ],
    right_align_cols=[1]
)

# --------------------------------------------------------------
# 1️⃣3️⃣  VISUAL REPORT SECTION – ORDERED
# --------------------------------------------------------------
doc.add_page_break()
doc.add_heading("Visual Report Section – Allocation & Diversification", level=1)

# 1) Ticker-level allocation
doc.add_heading("Ticker-Level Allocation Breakdown", level=2)
doc.add_picture(ticker_pie_stream, width=Inches(5.5))
p = doc.add_paragraph("Figure 1: Allocation by ticker — Core only (excludes AMZN & NDAQ).")
p.alignment = WD_ALIGN_PARAGRAPH.CENTER

# 2) Asset-class allocation
doc.add_paragraph("\n")
doc.add_heading("Asset Class Allocation Breakdown", level=2)
doc.add_picture(asset_pie_stream, width=Inches(5.5))
p = doc.add_paragraph("Figure 2: Allocation by asset class — Core only (excludes AMZN & NDAQ).")
p.alignment = WD_ALIGN_PARAGRAPH.CENTER

# 3) Sector heatmap
doc.add_paragraph("\n")
doc.add_heading("Sector Allocation Heatmap", level=2)
doc.add_picture(sector_stream, width=Inches(5.5))
p = doc.add_paragraph("Figure 3: Approximate sector exposure based on underlying holdings.")
p.alignment = WD_ALIGN_PARAGRAPH.CENTER


# --------------------------------------------------------------
# Performance & Growth
# --------------------------------------------------------------
doc.add_page_break()
doc.add_heading("Performance & Growth", level=1)

# 5) Performance vs benchmarks (MTD & YTD) — tables instead of chart
doc.add_heading("Performance vs Benchmarks (MTD & YTD)", level=2)

# Make sure MTD/YTD are numeric
bench_df["MTD %"] = pd.to_numeric(bench_df["MTD %"], errors="coerce")
bench_df["YTD %"] = pd.to_numeric(bench_df["YTD %"], errors="coerce")

# Get portfolio row
port_row = bench_df.loc[bench_df["Benchmark"] == "Portfolio (Live)"].iloc[0]

# Build MTD rows: one per benchmark (excluding portfolio)
mtd_rows = []
for _, row in bench_df.iterrows():
    if row["Benchmark"] == "Portfolio (Live)":
        continue
    bm = row["Benchmark"]
    port_mtd = port_row["MTD %"]
    bm_mtd = row["MTD %"]
    excess_mtd = port_mtd - bm_mtd if not pd.isna(port_mtd) and not pd.isna(bm_mtd) else None
    mtd_rows.append([
        bm,
        fmt_pct(port_mtd),
        fmt_pct(bm_mtd),
        fmt_pct(excess_mtd),
    ])

doc.add_paragraph("Month-to-date (MTD) comparison:")
add_table(
    ["Benchmark", "Portfolio MTD %", "Benchmark MTD %", "Excess MTD %"],
    mtd_rows,
    right_align_cols=[1, 2, 3],
)

doc.add_paragraph()  # blank line

# Build YTD rows
ytd_rows = []
for _, row in bench_df.iterrows():
    if row["Benchmark"] == "Portfolio (Live)":
        continue
    bm = row["Benchmark"]
    port_ytd = port_row["YTD %"]
    bm_ytd = row["YTD %"]
    excess_ytd = port_ytd - bm_ytd if not pd.isna(port_ytd) and not pd.isna(bm_ytd) else None
    ytd_rows.append([
        bm,
        fmt_pct(port_ytd),
        fmt_pct(bm_ytd),
        fmt_pct(excess_ytd),
    ])

doc.add_paragraph("Year-to-date (YTD) comparison:")
add_table(
    ["Benchmark", "Portfolio YTD %", "Benchmark YTD %", "Excess YTD %"],
    ytd_rows,
    right_align_cols=[1, 2, 3],
)

doc.add_paragraph()
doc.add_heading("Holdings Multi-Horizon Returns", level=2)
doc.add_paragraph(
    "Performance of each holding over multiple lookback periods, ranked by 6-month return:"
)

# Sort by 6-month return (descending)
returns_sorted = returns_df.sort_values("6M %", ascending=False, na_position="last")

rows = []
for _, r in returns_sorted.iterrows():
    rows.append([
        r["Ticker"],
        fmt_pct(r["1W %"]),
        fmt_pct(r["1M %"]),
        fmt_pct(r["3M %"]),
        fmt_pct(r["6M %"]),
    ])

add_table(
    ["Ticker", "1W %", "1M %", "3M %", "6M %"],
    rows,
    right_align_cols=[1, 2, 3, 4]
)


# 6) Compound value breakdown
doc.add_paragraph("\n")
doc.add_heading("Compound Value Breakdown", level=2)
doc.add_picture(compound_stream, width=Inches(6))
p = doc.add_paragraph(
    "Figure 6: Growth of your portfolio over time, separating your contributions from investment gains."
)
p.alignment = WD_ALIGN_PARAGRAPH.CENTER

# 7) 20-year projection
doc.add_paragraph("\n")
doc.add_heading("20-Year Projection Scenarios", level=2)

proj_headers = ["Year", "5% (Lump)", "7% (Lump)", "9% (Lump)",
                "5% (+$500/mo)", "7% (+$500/mo)", "9% (+$500/mo)"]
proj_table = doc.add_table(rows=1, cols=len(proj_headers))
proj_table.style = "Light Grid Accent 1"
hdr_cells = proj_table.rows[0].cells
for i, h in enumerate(proj_headers):
    hdr_cells[i].text = h
    for p in hdr_cells[i].paragraphs:
        for r in p.runs:
            r.bold = True

for row in proj_rows:
    row_cells = proj_table.add_row().cells
    row_cells[0].text = str(row[0])
    for i in range(1, len(proj_headers)):
        row_cells[i].text = f"{row[i]:,.0f}"

for row in proj_table.rows[1:]:
    for i in range(1, len(proj_headers)):
        for p in row.cells[i].paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.RIGHT

doc.add_paragraph("\n")
doc.add_picture(growth_stream, width=Inches(6))
p = doc.add_paragraph("Figure 7: Long-term projections under different return and contribution assumptions.")
p.alignment = WD_ALIGN_PARAGRAPH.CENTER

# --------------------------------------------------------------
# Risk & Volatility Section
# --------------------------------------------------------------
doc.add_page_break()
doc.add_heading("Risk & Volatility Analysis", level=1)

doc.add_heading("Expected Volatility by Asset Class", level=2)
doc.add_picture(vol_stream, width=Inches(5.5))
p = doc.add_paragraph("Figure 8: Approximate volatility (standard deviation) by asset class.")
p.alignment = WD_ALIGN_PARAGRAPH.CENTER

doc.add_paragraph("\n")
doc.add_heading("Risk vs Expected Return", level=2)
doc.add_picture(risk_stream, width=Inches(5.5))
p = doc.add_paragraph("Figure 9: Trade-off between expected return and volatility by asset class.")
p.alignment = WD_ALIGN_PARAGRAPH.CENTER

# --------------------------------------------------------------
# 1️⃣4️⃣  SAVE DOCX + PDF
# --------------------------------------------------------------
today_str = datetime.now().strftime("%Y-%m-%d")
docx_name = f"Investment_Report_{today_str}.docx"
pdf_name = f"Investment_Report_{today_str}.pdf"

doc.save(docx_name)

# Try to convert to PDF
try:
    convert(docx_name, pdf_name)
    print(f"Report generated: {docx_name}  and  {pdf_name}")
except Exception as e:
    print(f"Report generated: {docx_name}")
    print(f"PDF export failed: {e}")
    print("If you want automatic PDF export, install `docx2pdf` and ensure Word/LibreOffice is available.")

