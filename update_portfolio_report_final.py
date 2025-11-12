# ==============================================================
# update_portfolio_report_v3.py  (Regenerated)
# Live portfolio report with allocation, diversification,
# and real-time MTD / YTD benchmark comparisons
# ==============================================================

import yfinance as yf
import pandas as pd
import numpy as np
from io import BytesIO
import matplotlib.pyplot as plt
from datetime import datetime
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx2pdf import convert

# -------------------------- Formatting --------------------------

def fmt_pct(x):
    import math
    if x is None or (isinstance(x, float) and math.isnan(x)):
        return "-"
    return f"{x:+.2f}%"

def fmt_dollar(x):
    import math
    if x is None or (isinstance(x, float) and math.isnan(x)):
        return "-"
    return f"${x:,.2f}"

# ----------------------------- CONFIG -----------------------------

# Your holdings file in the repo
HOLDINGS_CSV = "sample holdings.csv"

# Optional: per-asset-class target file (asset_class,target_pct)
ASSET_TARGETS_CSV = "targets_asset.csv"   # if missing, it's fine

# How to split an asset-class target across its tickers:
#   "value" -> proportional to current market value (default)
#   "equal" -> equal weight among tickers in that class
TARGET_SPLIT_METHOD = "value"

RISK_FREE_RATE = 0.04
monthly_contrib = 500.0
COLOR_MAIN = ["#2563EB", "#10B981", "#F59E0B", "#6366F1", "#14B8A6"]

# Illustrative long-run assumptions for Risk/Return views
RISK_RETURN = {
    "US Equities":            {"return": 8.0,  "vol": 15.0},
    "International Equity":   {"return": 8.5,  "vol": 17.0},
    "Emerging Markets":       {"return": 9.0,  "vol": 20.0},
    "Fixed Income":           {"return": 4.0,  "vol": 5.0},
    "Real Estate":            {"return": 6.0,  "vol": 12.0},
    "Energy":                 {"return": 6.5,  "vol": 18.0},
    "Innovation/Tech":        {"return": 10.0, "vol": 25.0},
    "Commodities":            {"return": 6.0,  "vol": 10.0},
    "Gold / Precious Metals": {"return": 5.5,  "vol": 12.0},
}

# ----------------------- Helpers / Utilities -----------------------

def _norm_col(df: pd.DataFrame, want: str) -> str:
    want = want.strip().lower()
    for c in df.columns:
        if c.strip().lower() == want:
            return c
    raise ValueError(f"Required column '{want}' not found. Present: {list(df.columns)}")

def normalize_allocations(series: pd.Series) -> pd.Series:
    if len(series) == 0:
        return series
    rounded = series.round(2)
    diff = 100.00 - float(rounded.sum())
    rounded.iloc[-1] = round(float(rounded.iloc[-1]) + diff, 2)
    return rounded

def get_live_price(ticker: str) -> float:
    data = yf.Ticker(ticker).history(period="1d")
    if data.empty:
        data = yf.Ticker(ticker).history(period="5d")
        if data.empty:
            raise ValueError(f"No price data for {ticker}")
    return float(data["Close"].iloc[-1])

def get_return_pct(ticker, start_date, end_date):
    try:
        data = yf.download(ticker, start=start_date, end=end_date, progress=False)
        if data.empty:
            data = yf.download(ticker, period="60d", progress=False)
        if data.empty:
            return np.nan
        close = data["Close"]
        end = float(close.iloc[-1])
        start = float(close.iloc[0])
        return (end / start - 1.0) * 100.0
    except Exception:
        return np.nan

def read_asset_targets(path: str) -> pd.DataFrame | None:
    try:
        df = pd.read_csv(path)
    except Exception:
        return None
    try:
        ac = _norm_col(df, "asset_class")
        tp = _norm_col(df, "target_pct")
    except Exception:
        return None
    out = df.rename(columns={ac: "asset_class", tp: "target_pct"}).copy()
    out["asset_class"] = out["asset_class"].astype(str).str.strip()
    out["target_pct"] = pd.to_numeric(out["target_pct"], errors="coerce").fillna(0.0)
    # Normalize to 100 if not already
    if out["target_pct"].sum() > 0:
        out["target_pct"] = out["target_pct"] / out["target_pct"].sum() * 100.0
        out["target_pct"] = normalize_allocations(out["target_pct"])
    return out

def build_ticker_targets(df_holdings: pd.DataFrame,
                         df_asset_targets: pd.DataFrame | None,
                         split_method: str = "value") -> dict:
    """
    Priority:
      1) If holdings has per-ticker 'target_pct', use those (normalized).
      2) Else if df_asset_targets provided, distribute to tickers within each asset_class.
      3) Else equal-weight all tickers.
    """
    cols_l = [c.lower().strip() for c in df_holdings.columns]
    tickers = df_holdings["ticker"]
    # 1) Per-ticker targets in holdings
    if "target_pct" in cols_l:
        tcol = df_holdings.columns[cols_l.index("target_pct")]
        tgt = pd.to_numeric(df_holdings[tcol], errors="coerce").fillna(0.0)
        if tgt.sum() > 0:
            scaled = tgt / tgt.sum() * 100.0
            scaled = normalize_allocations(scaled.reset_index(drop=True))
            return dict(zip(tickers, scaled.tolist()))
    # 2) Asset-class targets
    if df_asset_targets is not None and not df_asset_targets.empty:
        merged = df_holdings.merge(
            df_asset_targets, on="asset_class", how="left", suffixes=("", "_ac")
        )
        targets = {}
        for ac, chunk in merged.groupby("asset_class", dropna=False):
            ac_target = float(chunk["target_pct"].iloc[0]) if "target_pct" in chunk else 0.0
            if ac_target <= 0 or len(chunk) == 0:
                continue
            if split_method == "equal":
                each = ac_target / len(chunk)
                for _, r in chunk.iterrows():
                    targets[r["ticker"]] = targets.get(r["ticker"], 0.0) + each
            else:  # proportional to current value
                vals = chunk["value"].clip(lower=0.0)
                denom = float(vals.sum())
                if denom == 0:
                    each = ac_target / len(chunk)
                    for _, r in chunk.iterrows():
                        targets[r["ticker"]] = targets.get(r["ticker"], 0.0) + each
                else:
                    for _, r in chunk.iterrows():
                        w = float(r["value"]) / denom if denom > 0 else 1.0 / len(chunk)
                        targets[r["ticker"]] = targets.get(r["ticker"], 0.0) + (ac_target * w)
        if targets:
            s = sum(targets.values())
            if s > 0:
                for k in list(targets.keys()):
                    targets[k] = targets[k] / s * 100.0
            keys = list(targets.keys())
            vals = [round(v, 2) for v in targets.values()]
            diff = round(100.0 - sum(vals), 2)
            vals[-1] = round(vals[-1] + diff, 2)
            return dict(zip(keys, vals))
    # 3) Equal-weight
    n = len(df_holdings)
    if n == 0:
        return {}
    eq = [round(100.0 / n, 2)] * n
    eq[-1] = round(eq[-1] + (100.0 - sum(eq)), 2)
    return dict(zip(tickers, eq))

# --------------------- 1) LOAD HOLDINGS + PRICES ---------------------

raw = pd.read_csv(HOLDINGS_CSV)

tcol = _norm_col(raw, "ticker")
scol = _norm_col(raw, "shares")

df = raw.rename(columns={tcol: "ticker", scol: "shares"}).copy()
df["ticker"] = df["ticker"].astype(str).str.upper().str.strip()
df["shares"] = pd.to_numeric(df["shares"], errors="coerce").fillna(0.0)

# asset_class optional
if "asset_class" not in [c.strip().lower() for c in df.columns]:
    df["asset_class"] = "Unknown"
else:
    ac_col = [c for c in df.columns if c.strip().lower() == "asset_class"][0]
    df.rename(columns={ac_col: "asset_class"}, inplace=True)

# Fetch prices; drop tickers that fail
prices = []
for t in df["ticker"]:
    try:
        prices.append(get_live_price(t))
    except Exception:
        prices.append(np.nan)
df["price"] = prices
df = df.dropna(subset=["price"]).reset_index(drop=True)

df["value"] = df["shares"] * df["price"]
total_value = float(df["value"].sum()) if len(df) else 0.0
df["allocation_pct"] = np.where(total_value > 0, (df["value"] / total_value * 100.0), 0.0)
df["allocation_pct"] = normalize_allocations(df["allocation_pct"])

# ------------- 2) ASSET-CLASS SUMMARY (entire portfolio) -------------

asset_df = df.groupby("asset_class", as_index=False)["value"].sum()
asset_df["allocation_pct"] = np.where(total_value > 0, (asset_df["value"] / total_value * 100.0), 0.0)
asset_df["allocation_pct"] = normalize_allocations(asset_df["allocation_pct"])

# ------------------ TARGETS (user-determined) -------------------
asset_targets_df = read_asset_targets(ASSET_TARGETS_CSV)
# build_ticker_targets will use per-ticker target_pct from HOLDINGS_CSV if present,
# otherwise it will use ASSET_TARGETS_CSV (by asset_class), else equal-weight.
TICKER_TARGETS_PCT = build_ticker_targets(df.copy(), asset_targets_df, TARGET_SPLIT_METHOD)

core_df = df.copy()
core_total = total_value
core_df["core_allocation_pct"] = np.where(core_total > 0, (core_df["value"] / core_total * 100.0), 0.0)

# Use a single 'target_pct' column (avoid target_pct_x / _y suffixes)
core_df["target_pct"] = core_df["ticker"].map(TICKER_TARGETS_PCT).astype(float).fillna(0.0)
core_df["target_value"] = core_total * (core_df["target_pct"] / 100.0)

# --- prevent suffix collisions if df already had these columns ---
for col in ["target_pct", "target_value", "core_allocation_pct"]:
    if col in df.columns:
        df.drop(columns=[col], inplace=True)

# merge cleanly with stable names
df = df.merge(
    core_df[["ticker", "target_value", "target_pct", "core_allocation_pct"]],
    on="ticker",
    how="left",
)

df["delta_to_target_raw"] = df["target_value"] - df["value"]
df["contribute_to_target"] = np.where(
    df["delta_to_target_raw"].isna(), np.nan,
    np.where(df["delta_to_target_raw"] > 0, df["delta_to_target_raw"], 0.0)
)


# -------- 3) SECTOR & GEOGRAPHIC (simple illustrative weights) --------

sector_weights = {
    "Information Technology": 32,
    "Financials": 12,
    "Healthcare": 11,
    "Industrials": 9,
    "Consumer Discretionary": 10,
    "Communication Services": 8,
    "Energy": 4,
    "Materials": 3,
    "Real Estate": 3,
    "Utilities": 3,
    "Bitcoin / Digital Assets": 3,
    "Gold / Precious Metals": 2,
}
sector_df_static = pd.DataFrame(
    {"Sector": list(sector_weights.keys()), "Weight": list(sector_weights.values())}
)

geo_alloc = {
    "United States": 70,
    "International Developed": 20,
    "Emerging Markets": 5,
    "Global Bonds": 3,
    "Precious Metals": 2,
}
geo_df = pd.DataFrame({"Region": list(geo_alloc.keys()), "Weight": list(geo_alloc.values())})

# ---------------------- 4) SECTOR HEATMAP CHART ----------------------

plt.figure(figsize=(6, 5))
plt.barh(
    sector_df_static["Sector"], sector_df_static["Weight"],
    color=plt.cm.Blues(np.linspace(0.4, 0.9, len(sector_df_static)))
)
plt.xlabel("Portfolio Exposure (%)")
plt.title("Sector Allocation Heatmap", fontsize=12, weight="bold")
plt.gca().invert_yaxis()
plt.tight_layout()
sector_stream = BytesIO()
plt.savefig(sector_stream, format="png", bbox_inches="tight", facecolor="white")
sector_stream.seek(0)
plt.close()

# -------------------- 5) BENCHMARK DATA (MTD & YTD) -------------------

benchmarks = {
    "S&P 500": "^GSPC",
    "Global 60/40": "AOR",
    "Conservative 40/60": "AOK",
}
today = datetime.now()
start_month = today.replace(day=1) - pd.tseries.offsets.BDay(1)
start_year = datetime(today.year, 1, 1) - pd.tseries.offsets.BDay(1)

bench_rows = []
for name, ticker in benchmarks.items():
    mtd = get_return_pct(ticker, start_month, today)
    ytd = get_return_pct(ticker, start_year, today)
    bench_rows.append(
        {"Benchmark": name,
         "MTD %": round(mtd, 2) if not np.isnan(mtd) else np.nan,
         "YTD %": round(ytd, 2) if not np.isnan(ytd) else np.nan}
    )

# Portfolio MTD/YTD as allocation-weighted average of its holdings
port_changes_mtd, port_changes_ytd, weights = [], [], []
for _, row in df.iterrows():
    t = row["ticker"]; w = row["allocation_pct"]
    mtd = get_return_pct(t, start_month, today)
    ytd = get_return_pct(t, start_year, today)
    if np.isnan(mtd) or np.isnan(ytd):
        continue
    port_changes_mtd.append(mtd); port_changes_ytd.append(ytd); weights.append(w)

def weighted_avg(values, weights):
    if not values:
        return np.nan
    v = np.array(values, dtype=float); w = np.array(weights, dtype=float)
    return float((v * w).sum() / w.sum()) if w.sum() != 0 else np.nan

port_mtd = weighted_avg(port_changes_mtd, weights)
port_ytd = weighted_avg(port_changes_ytd, weights)
bench_rows.insert(0, {"Benchmark": "Portfolio (Live)",
                      "MTD %": round(port_mtd, 2) if not np.isnan(port_mtd) else np.nan,
                      "YTD %": round(port_ytd, 2) if not np.isnan(port_ytd) else np.nan})
bench_df = pd.DataFrame(bench_rows)

# --------- 6) HOLDINGS MULTI-HORIZON RETURNS (1W, 1M, 3M, 6M) ---------

horizons = {"1W %": 7, "1M %": 30, "3M %": 90, "6M %": 180}
returns_rows = []
for _, row in df.iterrows():
    t = row["ticker"]
    r = {"Ticker": t}
    for label, days in horizons.items():
        start = today - pd.Timedelta(days=days)
        val = get_return_pct(t, start, today)
        r[label] = round(val, 2) if not np.isnan(val) else np.nan
    returns_rows.append(r)
returns_df = pd.DataFrame(returns_rows)

# ---------------- 7) LONG-TERM PROJECTIONS (20 YEARS) -----------------

rates = [0.05, 0.07, 0.09]
years = [1, 5, 10, 15, 20]

def future_value(principal, rate, years):
    return principal * ((1 + rate/12) ** (years * 12))

def future_value_with_contrib(principal, rate, years, monthly):
    months = years * 12
    monthly_rate = rate / 12
    fv_principal = principal * ((1 + monthly_rate) ** months)
    fv_contrib = monthly * (((1 + monthly_rate) ** months - 1) / monthly_rate)
    return fv_principal + fv_contrib

proj_rows = []
for y in years:
    row = [y]
    for r in rates:
        row.append(round(future_value(total_value, r, y)))
    for r in rates:
        row.append(round(future_value_with_contrib(total_value, r, y, monthly_contrib)))
    proj_rows.append(row)

plt.figure(figsize=(6, 4))
for i, r in enumerate(rates):
    plt.plot(years, [future_value(total_value, r, y) for y in years], label=f"{int(r*100)}% Lump Sum", linestyle='-')
    plt.plot(years, [future_value_with_contrib(total_value, r, y, monthly_contrib) for y in years], label=f"{int(r*100)}% + ${int(monthly_contrib)}/mo", linestyle='--')
plt.title("Portfolio Growth Projections (20-Year Scenarios)", fontsize=12, weight="bold")
plt.xlabel("Years"); plt.ylabel("Portfolio Value ($)")
plt.grid(alpha=0.3); plt.legend(fontsize=8)
plt.tight_layout()
growth_stream = BytesIO()
plt.savefig(growth_stream, format="png", bbox_inches="tight", facecolor="white")
growth_stream.seek(0)
plt.close()

# -------- 8) COMPOUND VALUE BREAKDOWN (CONTRIB VS GROWTH) -------------

years_compound = list(range(0, 21))
rate_for_compound = 0.07
contrib_values, growth_values = [], []
for y in years_compound:
    total_with_contrib = future_value_with_contrib(total_value, rate_for_compound, y, monthly_contrib)
    total_contrib = monthly_contrib * 12 * y
    contrib_values.append(total_contrib)
    growth_values.append(max(total_with_contrib - total_contrib, 0))
plt.figure(figsize=(6, 4))
plt.stackplot(years_compound, contrib_values, growth_values, labels=["Contributions", "Growth"], colors=[COLOR_MAIN[0], COLOR_MAIN[1]], alpha=0.85)
plt.title("Compound Value Breakdown (Contributions vs Growth)", fontsize=12, weight="bold")
plt.xlabel("Years"); plt.ylabel("Portfolio Value ($)")
plt.legend(loc="upper left", fontsize=8)
plt.tight_layout()
compound_stream = BytesIO()
plt.savefig(compound_stream, format="png", bbox_inches="tight", facecolor="white")
compound_stream.seek(0)
plt.close()

# -------------- 9) PERFORMANCE VS BENCHMARKS (CHART) ------------------

bench_df["MTD %"] = pd.to_numeric(bench_df["MTD %"], errors="coerce")
bench_df["YTD %"] = pd.to_numeric(bench_df["YTD %"], errors="coerce")
bench_plot = bench_df.dropna(subset=["MTD %", "YTD %"], how="all").reset_index(drop=True)

if bench_plot.empty:
    plt.figure(figsize=(6, 4))
    plt.text(0.5, 0.5, "Benchmark return data unavailable.\n(Check connection or market days.)", ha="center", va="center", fontsize=9)
    plt.axis("off")
    plt.tight_layout()
else:
    x = np.arange(len(bench_plot))
    width = 0.35
    mtd_vals = bench_plot["MTD %"].to_numpy(dtype=float)
    ytd_vals = bench_plot["YTD %"].to_numpy(dtype=float)

    plt.figure(figsize=(6, 4))
    mtd_bars = plt.bar(x - width/2, mtd_vals, width, label="MTD", color=COLOR_MAIN[0])
    ytd_bars = plt.bar(x + width/2, ytd_vals, width, label="YTD", color=COLOR_MAIN[1])
    plt.xticks(x, bench_plot["Benchmark"], rotation=20, ha="right")
    plt.ylabel("Return (%)")
    plt.title("Portfolio vs Benchmarks (MTD & YTD)", fontsize=12, weight="bold")
    plt.legend(fontsize=8)
    plt.grid(axis="y", alpha=0.3)

    vals = np.concatenate([mtd_vals[~np.isnan(mtd_vals)], ytd_vals[~np.isnan(ytd_vals)]]) if (len(mtd_vals) + len(ytd_vals)) else np.array([])
    if vals.size > 0:
        vmin, vmax = float(np.min(vals)), float(np.max(vals))
        if vmin == vmax:
            vmin -= 1.0; vmax += 1.0
        else:
            pad = (vmax - vmin) * 0.2
            vmin -= pad; vmax += pad
        plt.ylim(vmin, vmax)

    def label_bars(bars, vals):
        for bar, val in zip(bars, vals):
            if np.isnan(val): continue
            height = bar.get_height()
            plt.text(bar.get_x() + bar.get_width()/2, height, f"{val:+.1f}%", ha="center", va="bottom", fontsize=7)

    label_bars(mtd_bars, mtd_vals); label_bars(ytd_bars, ytd_vals)
    plt.tight_layout()

# ------------------ 10) RISK & VOLATILITY (CLEANED) ------------------

present_assets = sorted(set(asset_df["asset_class"]))
risk_rows = []
for ac in present_assets:
    base = RISK_RETURN.get(ac)
    if base is None:
        if ac.lower().startswith("international"): base = RISK_RETURN.get("International Equity")
        elif ac.lower().startswith("emerging"):     base = RISK_RETURN.get("Emerging Markets")
        elif "gold" in ac.lower() or "precious" in ac.lower(): base = RISK_RETURN.get("Gold / Precious Metals")
        elif "fixed" in ac.lower() or "bond" in ac.lower(): base = RISK_RETURN.get("Fixed Income")
        elif "real" in ac.lower(): base = RISK_RETURN.get("Real Estate")
        elif "energy" in ac.lower(): base = RISK_RETURN.get("Energy")
        elif "innovation" in ac.lower() or "tech" in ac.lower(): base = RISK_RETURN.get("Innovation/Tech")
        elif "commodit" in ac.lower(): base = RISK_RETURN.get("Commodities")
    if base:
        risk_rows.append({"asset_class": ac, "vol": base["vol"], "ret": base["return"]})

risk_df = pd.DataFrame(risk_rows)

# Volatility bar chart
plt.figure(figsize=(6, 4))
plt.bar(risk_df["asset_class"], risk_df["vol"], color=COLOR_MAIN[0])
plt.title("Expected Volatility by Asset Class", fontsize=12, weight="bold")
plt.ylabel("Standard Deviation (%)")
plt.xticks(rotation=25, ha="right")
if len(risk_df) and risk_df["vol"].max() == risk_df["vol"].min():
    ymin = risk_df["vol"].min() - 1
    ymax = risk_df["vol"].max() + 1
else:
    ymin = max(0, float(risk_df["vol"].min()) - 2)
    ymax = float(risk_df["vol"].max()) + 4
plt.ylim(ymin, ymax)
plt.grid(axis="y", alpha=0.25)
plt.tight_layout()
vol_stream = BytesIO()
plt.savefig(vol_stream, format="png", bbox_inches="tight", facecolor="white")
vol_stream.seek(0)
plt.close()

# Risk vs Return scatter
plt.figure(figsize=(6, 4))
plt.scatter(risk_df["vol"], risk_df["ret"], s=80, edgecolors="black", linewidths=0.6, alpha=0.9, color=COLOR_MAIN[1])
for _, row in risk_df.iterrows():
    plt.annotate(row["asset_class"], (row["vol"] + 0.4, row["ret"]), fontsize=8, weight="bold")
plt.title("Risk vs Expected Return by Asset Class", fontsize=12, weight="bold")
plt.xlabel("Volatility (Std Dev %)"); plt.ylabel("Expected Annual Return (%)")
if len(risk_df):
    xpad = (risk_df["vol"].max() - risk_df["vol"].min()) * 0.2 if risk_df["vol"].max() != risk_df["vol"].min() else 3
    ypad = (risk_df["ret"].max() - risk_df["ret"].min()) * 0.2 if risk_df["ret"].max() != risk_df["ret"].min() else 2
    plt.xlim(max(0, risk_df["vol"].min() - xpad), risk_df["vol"].max() + xpad)
    plt.ylim(max(0, risk_df["ret"].min() - ypad), risk_df["ret"].max() + ypad)
plt.grid(alpha=0.3)
plt.tight_layout()
risk_stream = BytesIO()
plt.savefig(risk_stream, format="png", bbox_inches="tight", facecolor="white")
risk_stream.seek(0)
plt.close()

# ---------------------- 11) ALLOCATION PIE CHARTS ----------------------

plt.figure(figsize=(6, 6))
plt.pie(
    df["allocation_pct"], labels=df["ticker"],
    autopct="%1.2f%%", startangle=90, pctdistance=0.85, labeldistance=1.05
)
plt.title("Portfolio Allocation by Ticker", fontsize=12, weight="bold")
plt.tight_layout()
ticker_pie_stream = BytesIO()
plt.savefig(ticker_pie_stream, format="png", bbox_inches="tight", facecolor="white")
ticker_pie_stream.seek(0)
plt.close()

plt.figure(figsize=(6, 6))
plt.pie(
    asset_df["allocation_pct"], labels=asset_df["asset_class"],
    autopct="%1.2f%%", startangle=90, pctdistance=0.85, labeldistance=1.05
)
plt.title("Asset Class Allocation", fontsize=12, weight="bold")
plt.tight_layout()
asset_pie_stream = BytesIO()
plt.savefig(asset_pie_stream, format="png", bbox_inches="tight", facecolor="white")
asset_pie_stream.seek(0)
plt.close()

# ------------------------- 12) BUILD WORD DOC --------------------------

doc = Document()
style = doc.styles["Normal"]
style.font.name = "Calibri"
style._element.rPr.rFonts.set(qn("w:eastAsia"), "Calibri")
style.font.size = Pt(11)

def add_table(headers, rows, right_align_cols=None):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Light Grid Accent 1"
    hdr_cells = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr_cells[i].text = h
        for p in hdr_cells[i].paragraphs:
            for r in p.runs:
                r.bold = True
    for row_data in rows:
        row_cells = table.add_row().cells
        for i, val in enumerate(row_data):
            row_cells[i].text = str(val)
    if right_align_cols is not None:
        for row in table.rows[1:]:
            for idx in right_align_cols:
                cell = row.cells[idx]
                for p in cell.paragraphs:
                    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    return table

# Cover
doc.add_paragraph("\n\n\n\n\n")
cover = doc.add_paragraph("Comprehensive Investment Report — Live Portfolio")
cover.runs[0].bold = True
cover.runs[0].font.size = Pt(26)
cover.alignment = WD_ALIGN_PARAGRAPH.CENTER

subtitle = doc.add_paragraph("Automated Summary, Diversification, and Performance vs Benchmarks")
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
subtitle.runs[0].font.size = Pt(14)
subtitle.runs[0].font.color.rgb = RGBColor(70, 130, 180)

doc.add_paragraph()
doc.add_paragraph("Prepared for: Your Mom").alignment = WD_ALIGN_PARAGRAPH.CENTER
doc.add_page_break()

# Executive Summary
doc.add_heading("Executive Summary", level=1)
doc.add_paragraph(
    f"This report summarizes your current portfolio based on live market data. "
    f"Total portfolio value is approximately ${total_value:,.0f}. "
    f"Targets can be supplied per-ticker or per-asset-class via CSV."
)

# Overall Summary
doc.add_heading("Overall Summary", level=1)
overall_rows = [["Total Portfolio Value", f"{total_value:,.2f}", "Sum of all holdings at live prices"]]
add_table(["Category", "Amount ($)", "Notes"], overall_rows, right_align_cols=[1])

# Holdings by Ticker
doc.add_heading("Holdings by Ticker", level=1)
ticker_rows = []
for _, row in df.iterrows():
    ticker_rows.append([
        row["ticker"],
        row["asset_class"],
        f"{row['shares']:.4f}",
        f"{row['price']:.2f}",
        f"{row['value']:.2f}",
        f"{row['core_allocation_pct']:.2f}%",
        f"{(row.get('target_pct') or 0.0):.2f}%",
        "-" if pd.isna(row.get("contribute_to_target")) else fmt_dollar(row["contribute_to_target"]),
    ])
add_table(
    ["Ticker", "Asset Class", "Shares", "Price ($)", "Value ($)", "Allocation %", "Target %", "Contribute to Target ($)"],
    ticker_rows,
    right_align_cols=[2, 3, 4, 5, 6, 7]
)

# Asset Class Allocation Overview
doc.add_heading("Asset Class Allocation Overview", level=1)
ac_rows = []
for _, row in asset_df.iterrows():
    ac_rows.append([row["asset_class"], f"{row['value']:.2f}", f"{row['allocation_pct']:.2f}%"])
add_table(["Asset Class", "Value ($)", "Allocation %"], ac_rows, right_align_cols=[1, 2])

# Next Steps & Ongoing Strategy
doc.add_page_break()
doc.add_heading("Next Steps & Ongoing Strategy", level=1)
add_table(
    ["Focus Area", "Guidance"],
    [
        ["Rebalancing", "Review annually; adjust if allocations drift ±5%."],
        ["Contributions", f"Automate around ${monthly_contrib:,.0f}/mo into your target allocation."],
        ["Risk Management", "Stay within your risk tolerance; increase bonds as you approach major goals."],
        ["Monitoring", "Avoid overreacting to short-term volatility; review quarterly or annually."],
    ],
)

# Totals Check
doc.add_heading("Totals Check", level=1)
add_table(["Metric", "Value"], [["Total Portfolio Value", f"{total_value:,.2f}"]], right_align_cols=[1])

# Visuals – Allocation & Diversification
doc.add_page_break()
doc.add_heading("Visual Report – Allocation & Diversification", level=1)
doc.add_heading("Ticker-Level Allocation Breakdown", level=2)
doc.add_picture(ticker_pie_stream, width=Inches(5.5))
p = doc.add_paragraph("Figure 1: Allocation by ticker.")
p.alignment = WD_ALIGN_PARAGRAPH.CENTER

doc.add_paragraph("\n")
doc.add_heading("Asset Class Allocation Breakdown", level=2)
doc.add_picture(asset_pie_stream, width=Inches(5.5))
p = doc.add_paragraph("Figure 2: Allocation by asset class.")
p.alignment = WD_ALIGN_PARAGRAPH.CENTER

doc.add_paragraph("\n")
doc.add_heading("Sector Allocation Heatmap", level=2)
doc.add_picture(sector_stream, width=Inches(5.5))
p = doc.add_paragraph("Figure 3: Approximate sector exposure based on underlying holdings.")
p.alignment = WD_ALIGN_PARAGRAPH.CENTER

# Performance & Growth
doc.add_page_break()
doc.add_heading("Performance & Growth", level=1)
doc.add_heading("Performance vs Benchmarks (MTD & YTD)", level=2)

bench_df["MTD %"] = pd.to_numeric(bench_df["MTD %"], errors="coerce")
bench_df["YTD %"] = pd.to_numeric(bench_df["YTD %"], errors="coerce")
if "Portfolio (Live)" in bench_df["Benchmark"].values:
    port_row = bench_df.loc[bench_df["Benchmark"] == "Portfolio (Live)"].iloc[0]
else:
    port_row = pd.Series({"MTD %": np.nan, "YTD %": np.nan})

mtd_rows, ytd_rows = [], []
for _, row in bench_df.iterrows():
    if row["Benchmark"] == "Portfolio (Live)":
        continue
    mtd_rows.append([row["Benchmark"], fmt_pct(port_row.get("MTD %")), fmt_pct(row["MTD %"]), fmt_pct((port_row.get("MTD %") or np.nan) - (row["MTD %"] or np.nan))])
    ytd_rows.append([row["Benchmark"], fmt_pct(port_row.get("YTD %")), fmt_pct(row["YTD %"]), fmt_pct((port_row.get("YTD %") or np.nan) - (row["YTD %"] or np.nan))])

doc.add_paragraph("Month-to-date (MTD) comparison:")
add_table(["Benchmark", "Portfolio MTD %", "Benchmark MTD %", "Excess MTD %"], mtd_rows, right_align_cols=[1, 2, 3])

doc.add_paragraph("Year-to-date (YTD) comparison:")
add_table(["Benchmark", "Portfolio YTD %", "Benchmark YTD %", "Excess YTD %"], ytd_rows, right_align_cols=[1, 2, 3])

doc.add_paragraph()
doc.add_heading("Holdings Multi-Horizon Returns", level=2)
doc.add_paragraph("Performance of each holding over multiple lookback periods, ranked by 6-month return:")
returns_sorted = returns_df.sort_values("6M %", ascending=False, na_position="last")
rows = []
for _, r in returns_sorted.iterrows():
    rows.append([r["Ticker"], fmt_pct(r["1W %"]), fmt_pct(r["1M %"]), fmt_pct(r["3M %"]), fmt_pct(r["6M %"])])
add_table(["Ticker", "1W %", "1M %", "3M %", "6M %"], rows, right_align_cols=[1, 2, 3, 4])

doc.add_paragraph("\n")
doc.add_heading("Compound Value Breakdown", level=2)
doc.add_picture(compound_stream, width=Inches(6))
p = doc.add_paragraph("Figure 4: Growth of your portfolio over time, separating contributions from investment gains.")
p.alignment = WD_ALIGN_PARAGRAPH.CENTER

doc.add_paragraph("\n")
doc.add_heading("20-Year Projection Scenarios", level=2)
proj_headers = ["Year", "5% (Lump)", "7% (Lump)", "9% (Lump)", "5% (+$500/mo)", "7% (+$500/mo)", "9% (+$500/mo)"]
proj_table = doc.add_table(rows=1, cols=len(proj_headers))
proj_table.style = "Light Grid Accent 1"
hdr_cells = proj_table.rows[0].cells
for i, h in enumerate(proj_headers):
    hdr_cells[i].text = h
    for p in hdr_cells[i].paragraphs:
        for r in p.runs:
            r.bold = True
for row in proj_rows:
    row_cells = proj_table.add_row().cells
    row_cells[0].text = str(row[0])
    for i in range(1, len(proj_headers)):
        row_cells[i].text = f"{row[i]:,.0f}"
for row in proj_table.rows[1:]:
    for i in range(1, len(proj_headers)):
        for p in row.cells[i].paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
doc.add_paragraph("\n")
doc.add_picture(growth_stream, width=Inches(6))
p = doc.add_paragraph("Figure 5: Long-term projections under different return and contribution assumptions.")
p.alignment = WD_ALIGN_PARAGRAPH.CENTER

# Risk & Volatility
doc.add_page_break()
doc.add_heading("Risk & Volatility Analysis", level=1)
doc.add_heading("Expected Volatility by Asset Class", level=2)
doc.add_picture(vol_stream, width=Inches(5.5))
p = doc.add_paragraph("Figure 6: Approximate volatility (standard deviation) by asset class.")
p.alignment = WD_ALIGN_PARAGRAPH.CENTER

doc.add_paragraph("\n")
doc.add_heading("Risk vs Expected Return", level=2)
doc.add_picture(risk_stream, width=Inches(5.5))
p = doc.add_paragraph("Figure 7: Trade-off between expected return and volatility by asset class.")
p.alignment = WD_ALIGN_PARAGRAPH.CENTER

# ------------------------- 13) SAVE DOCX + PDF ------------------------

today_str = datetime.now().strftime("%Y-%m-%d")
docx_name = f"Investment_Report_{today_str}.docx"
pdf_name = f"Investment_Report_{today_str}.pdf"

doc.save(docx_name)

try:
    convert(docx_name, pdf_name)
    print(f"Report generated: {docx_name}  and  {pdf_name}")
except Exception as e:
    print(f"Report generated: {docx_name}")
    print(f"PDF export failed: {e}")
    print("If you want automatic PDF export, install `docx2pdf` and ensure Word/LibreOffice is available.")
